#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Pass 1: Chèn Chương 3 vào Decuong_OmniRAG (2).docx
- Backup → Decuong_OmniRAG_backup.docx
- Output → Decuong_OmniRAG_v2.docx (chỉ có Chương 3, kiểm tra trước)
- Update sai sót nhỏ trong Chương 2
"""

import shutil
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK

SRC = "docs/Decuong_OmniRAG (2).docx"
BAK = "docs/Decuong_OmniRAG_backup.docx"
OUT = "docs/Decuong_OmniRAG_v2.docx"

# ── Backup ──────────────────────────────────────────────────────────────────
shutil.copy2(SRC, BAK)
print(f"Backup → {BAK}")

doc = docx.Document(SRC)

# ── Tìm điểm chèn ────────────────────────────────────────────────────────────
danh_muc = None
for p in doc.paragraphs:
    if "DANH MỤC TÀI LIỆU THAM KHẢO" in p.text and p.style.name == "Heading 1":
        danh_muc = p
        break
assert danh_muc, "Không tìm thấy DANH MỤC TÀI LIỆU (Heading 1)"
print(f"Điểm chèn: '{danh_muc.text[:60]}'")

# ── Helpers ──────────────────────────────────────────────────────────────────
def ins(text, style="Normal"):
    """Chèn đoạn văn TRƯỚC danh_muc."""
    p = doc.add_paragraph(text, style)
    danh_muc._element.addprevious(p._element)
    return p

def pb():
    """Page break trước danh_muc."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    danh_muc._element.addprevious(p._element)

def bul(text):
    """Bullet point."""
    ins(f"•  {text}")

def placeholder(label):
    """Placeholder cho hình vẽ."""
    ins(f"[{label}]")

# ═══════════════════════════════════════════════════════════════════════════
# CHƯƠNG 3
# ═══════════════════════════════════════════════════════════════════════════
pb()
ins("CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ NỀN TẢNG OMNIRAG", "Heading 1")

# ── 3.1 ────────────────────────────────────────────────────────────────────
ins("3.1. Phân tích chức năng và luồng nghiệp vụ", "Heading 2")
ins(
    "Phần này trình bày phân tích chức năng cốt lõi của OmniRAG thông qua bốn luồng nghiệp vụ "
    "chính: quy trình xây dựng bot tự phục vụ (Self-build Bot), cơ chế cô lập dữ liệu đa "
    "khách hàng, luồng xử lý tin nhắn đa kênh, và cơ chế Domain Profile Registry cấu hình "
    "pipeline RAG chuyên biệt theo lĩnh vực ứng dụng."
)

# 3.1.1
ins("3.1.1. Quy trình Self-build bot: Bot Wizard, Domain Selector và Template", "Heading 3")
ins(
    "OmniRAG cho phép bất kỳ người dùng nào — không cần kiến thức kỹ thuật chuyên sâu — "
    "tự tạo và triển khai một trợ lý ảo RAG thông qua giao diện Bot Wizard nhiều bước. "
    "Quy trình được thiết kế theo mô hình wizard dẫn dắt (guided wizard) gồm bốn bước tuần tự:"
)
bul(
    "Bước 1 — Chọn Template: Người dùng lựa chọn từ thư viện template có sẵn "
    "(Customer Support, Education Assistant, Legal Advisor, Sales Agent...). Mỗi template "
    "đóng gói sẵn system prompt, cấu hình model và domain phù hợp, giảm thời gian setup "
    "từ hàng giờ xuống còn dưới 5 phút. API endpoint GET /api/v1/bot-templates cung cấp "
    "danh sách template, hỗ trợ lọc theo domain qua GET /api/v1/bot-templates/domains/{domain}."
)
bul(
    "Bước 2 — Chọn Domain: Người dùng xác định lĩnh vực ứng dụng (General / Education / "
    "Legal / Sales). Domain quyết định toàn bộ cấu hình RAG pipeline phía sau: chiến lược "
    "phân đoạn văn bản, kích thước chunk, số lượng tài liệu truy xuất (top_k), và việc "
    "có bật Knowledge Graph hay không. Giao diện hiển thị mô tả ngắn gọn và icon màu sắc "
    "đặc trưng cho mỗi domain để hỗ trợ người dùng ra quyết định nhanh."
)
bul(
    "Bước 3 — Cấu hình nâng cao: Người dùng tùy chỉnh các tham số của bot: welcome_message, "
    "system_prompt, nhiệt độ (temperature), giới hạn token (max_tokens), giọng điệu trả lời "
    "(tone_formality), và ngưỡng similarity (similarity_threshold). Giao diện BotConfigPage "
    "luôn hiển thị toàn bộ tham số nâng cao, không ẩn sau toggle, cho phép người dùng "
    "kiểm soát chi tiết ngay từ bước khởi tạo."
)
bul(
    "Bước 4 — Xem lại và tạo bot: Tóm tắt toàn bộ cấu hình, xác nhận và gửi yêu cầu "
    "POST /api/v1/bots. Backend tự động sinh api_key dài 32 ký tự URL-safe, tạo bản ghi "
    "trong PostgreSQL và trả về thông tin bot đầy đủ. Sau khi tạo thành công, người dùng "
    "được điều hướng sang trang cấu hình để upload tài liệu và tích hợp kênh."
)
ins(
    "Toàn bộ thông tin cấu hình bot — bao gồm tham số RAG, cấu hình Zalo, system prompt — "
    "được lưu trong cột config kiểu JSONB của bảng bots (PostgreSQL). Thiết kế JSONB cho phép "
    "mở rộng schema cấu hình mà không cần migration database, phù hợp với chu kỳ phát triển "
    "nhanh của sản phẩm SaaS."
)

# 3.1.2
ins("3.1.2. Cơ chế phân tách dữ liệu đa khách hàng (Isolated Multi-tenancy)", "Heading 3")
ins(
    "Trong môi trường SaaS đa khách hàng, rủi ro rò rỉ dữ liệu (data leakage) giữa các "
    "tenant là mối đe dọa nghiêm trọng nhất. OmniRAG triển khai kiến trúc Bridge Pattern "
    "(Fernandez et al., 2023) — kết hợp giữa Silo Pattern (cô lập cao) và Pool Pattern "
    "(hiệu quả chi phí) — thực thi cô lập dữ liệu qua năm tầng lưu trữ:"
)
bul(
    "Tầng quan hệ (PostgreSQL): Mọi bảng dữ liệu đều có cột tenant_id hoặc bot_id làm "
    "khóa ngoại. Truy vấn SQL luôn bao gồm điều kiện lọc theo tenant_id của người dùng "
    "hiện tại, được inject bởi dependency injection layer (get_current_tenant). Cascade "
    "delete đảm bảo khi xóa tenant, toàn bộ dữ liệu liên quan bị xóa sạch."
)
bul(
    "Tầng vector (Qdrant): Tất cả vector embedding được lưu trong một collection duy nhất "
    "(omnirag) nhưng mỗi document chunk mang payload {bot_id} làm định danh phân tách. "
    "Mọi truy vấn vector search đều đính kèm bộ lọc bot_id, đảm bảo bot A không bao giờ "
    "truy xuất được dữ liệu của bot B dù cùng chung hạ tầng."
)
bul(
    "Tầng đồ thị tri thức (LightRAG): Mỗi bot sở hữu workspace LightRAG độc lập thông qua "
    "hai tham số workspace=bot_id (cô lập trong Qdrant) và working_dir=./rag_storage/lightrag_{bot_id} "
    "(cô lập file system). Knowledge graph của bot giáo dục không thể bị truy xuất bởi "
    "bot pháp lý của tenant khác."
)
bul(
    "Tầng tài liệu phi cấu trúc (MongoDB): Các collection conversations, sessions, api_keys "
    "đều lưu trữ user_id/bot_id và áp dụng scoped query. Lịch sử hội thoại của khách hàng "
    "A hoàn toàn riêng biệt với khách hàng B ngay cả khi sử dụng cùng một bot template."
)
bul(
    "Tầng lưu trữ file (MinIO): File tài liệu được lưu với UUID-based object key "
    "(UUID ngẫu nhiên, không phụ thuộc tên file gốc), ngăn chặn path traversal attack. "
    "Celery worker download file theo object key sau khi xác minh quyền truy cập."
)
ins(
    "Theo khuyến nghị từ Multi-Tenant AI SaaS Architecture (2025), OmniRAG áp dụng thêm "
    "Row-Level Security (RLS) trên PostgreSQL và strict RBAC với ba vai trò: owner (toàn quyền), "
    "admin (quản lý bot và tài liệu), member (chỉ chat). Mô hình này đáp ứng yêu cầu của "
    "các tổ chức xử lý dữ liệu nhạy cảm như hồ sơ pháp lý và thông tin y tế."
)

# 3.1.3
ins("3.1.3. Luồng xử lý tin nhắn đa kênh (Omnichannel Message Flow)", "Heading 3")
ins(
    "OmniRAG hỗ trợ hai kênh tích hợp Zalo với kiến trúc và luồng xử lý khác nhau, "
    "được thiết kế để phù hợp với từng mô hình triển khai của doanh nghiệp:"
)
ins(
    "Kênh 1 — Zalo Bot Direct (bot-api.zapps.me): Tích hợp trực tiếp không qua trung gian, "
    "phù hợp cho doanh nghiệp muốn sở hữu toàn bộ luồng xử lý. Luồng hoạt động như sau:"
)
bul(
    "Bước kết nối: Admin gửi POST /api/v1/channels/zalo-bot/connect với bot_id và bot_token. "
    "Hệ thống gọi API getMe (xác minh token), sinh webhook_secret 24 ký tự ngẫu nhiên, "
    "gọi API setWebhook để đăng ký endpoint {PUBLIC_URL}/api/v1/channels/zalo-bot/webhook/{bot_id}. "
    "Thông tin kết nối (token, secret, trạng thái) được lưu vào bot.config.zalo_bot."
)
bul(
    "Bước nhận tin nhắn: Khi người dùng Zalo gửi tin, Zalo Platform gọi webhook của OmniRAG. "
    "Backend xác minh chữ ký HMAC-SHA256 qua header x-bot-api-secret-token bằng hàm "
    "hmac.compare_digest() (constant-time, chống timing attack). Nếu hợp lệ, "
    "dispatch Celery task process_zalo_bot_webhook_task bất đồng bộ."
)
bul(
    "Bước xử lý: Celery worker gọi handle_webhook() — trích xuất chat_id và nội dung tin, "
    "gửi tín hiệu 'typing' qua sendChatAction để người dùng biết bot đang xử lý, "
    "gọi RAG pipeline và trả về câu trả lời qua API sendMessage."
)
ins(
    "Kênh 2 — Zalo OA qua func.vn Hub: Tích hợp qua trung gian func.vn, phù hợp cho "
    "doanh nghiệp đang sử dụng Zalo Official Account (OA). Luồng hoạt động:"
)
bul(
    "Webhook tập trung: func.vn gửi tất cả tin nhắn từ mọi OA đến một endpoint duy nhất "
    "POST /api/v1/channels/zalo/hub-webhook với payload chứa account_id (ID của OA), "
    "sender_id (ID người dùng Zalo), và nội dung tin nhắn."
)
bul(
    "Định tuyến theo bot: Backend tra cứu bot có config.zalo_integration.account_id khớp "
    "với account_id nhận được — hỗ trợ cả định dạng có/không có tiền tố 'zu'. "
    "Sau khi tìm thấy bot phù hợp, gọi RAG pipeline và gửi phản hồi qua func.vn Reply API "
    "(POST đến FUNC_API_URL với Bearer token)."
)
ins(
    "Cả hai kênh đều áp dụng kiến trúc fire-and-forget với Celery: webhook handler phản hồi "
    "HTTP 200 ngay lập tức (tránh timeout từ Zalo Platform), trong khi xử lý RAG diễn ra "
    "bất đồng bộ ở background. Timeout tối đa cho RAG pipeline được giới hạn ở 60 giây."
)
placeholder(
    "Hình 3.1 — Omnichannel Message Flow: Sơ đồ sequence so sánh hai luồng "
    "Zalo Bot Direct và Zalo Hub, từ User → Platform → OmniRAG → RAG Engine → Response"
)

# 3.1.4
ins("3.1.4. Domain Profile Registry: cấu hình RAG chuyên biệt theo lĩnh vực", "Heading 3")
ins(
    "OmniRAG triển khai Domain Profile Registry — một cơ chế cấu hình tập trung "
    "cho phép pipeline RAG tự động điều chỉnh theo đặc thù của từng lĩnh vực ứng dụng. "
    "Mỗi domain profile định nghĩa một tập hợp các tham số tối ưu dựa trên đặc điểm "
    "ngôn ngữ và cấu trúc tài liệu của lĩnh vực đó. Bảng 3.1 tổng hợp cấu hình của "
    "bốn domain hiện có:"
)
# Tạo bảng 3.1 Domain Profile Registry
tbl = doc.add_table(rows=6, cols=7)
tbl.style = "Table Grid"

# Move table before danh_muc
danh_muc._element.addprevious(tbl._element)

headers = ["Domain", "Chiến lược\nchunking", "Chunk size\n(tokens)", "Overlap\n(tokens)", "Retrieval K\n(rerank top)", "LightRAG", "KG mode"]
data = [
    ["general", "recursive", "512", "64", "10 (top 5)", "Tắt", "naive"],
    ["education", "sentence", "384", "32", "12 (top 6)", "Bật", "local"],
    ["legal", "article", "1.024", "128", "8 (top 4)", "Bật", "hybrid"],
    ["sales", "recursive", "256", "32", "15 (top 7)", "Tắt", "naive"],
]

hdr_row = tbl.rows[0]
for i, h in enumerate(headers):
    hdr_row.cells[i].text = h

for row_idx, row_data in enumerate(data):
    row = tbl.rows[row_idx + 1]
    for col_idx, val in enumerate(row_data):
        row.cells[col_idx].text = val

ins("Bảng 3.1 — Domain Profile Registry trong OmniRAG (Nguồn: domain_config.py)")
danh_muc._element.addprevious(tbl._element)

ins(
    "Chiến lược article chunking (dùng cho domain Legal) là điểm khác biệt quan trọng: "
    "thay vì phân đoạn theo ký tự hay câu, hệ thống nhận dạng marker văn bản pháp lý "
    "tiếng Việt bằng biểu thức chính quy (?=Điều\\s+\\d+) để tách tại ranh giới điều luật. "
    "Điều này đảm bảo mỗi chunk giữ nguyên ý nghĩa pháp lý của một điều khoản, "
    "tránh hiện tượng điều luật bị cắt đứt ở giữa gây mất ngữ cảnh khi truy xuất."
)
ins(
    "Domain Education bật LightRAG ở chế độ local — tập trung truy vấn từ các entity node "
    "liên quan, phù hợp với câu hỏi chi tiết về khái niệm cụ thể (định nghĩa, ví dụ). "
    "Domain Legal bật LightRAG ở chế độ hybrid (kết hợp local + global), hỗ trợ cả "
    "câu hỏi về điều khoản cụ thể lẫn câu hỏi về mối quan hệ giữa các điều luật."
)

# ── 3.2 ────────────────────────────────────────────────────────────────────
ins("3.2. Thiết kế các biểu đồ hệ thống", "Heading 2")
ins(
    "Phần này trình bày các biểu đồ thiết kế hệ thống mô tả luồng xử lý nghiệp vụ "
    "và tương tác giữa các thành phần. Các biểu đồ được xây dựng theo chuẩn UML 2.5."
)

# 3.2.1
ins("3.2.1. Biểu đồ Use Case và Sequence Diagrams cho RAG Agent", "Heading 3")
ins(
    "Hệ thống OmniRAG xác định bốn nhóm actor chính tương tác với nền tảng:"
)
bul(
    "Admin (Tenant Owner): Đăng ký tenant, tạo/cấu hình bot, upload tài liệu, "
    "tích hợp kênh Zalo, xem analytics và quản lý thành viên trong tenant."
)
bul(
    "End User (Web): Người dùng nội bộ sử dụng giao diện Chat Playground qua trình duyệt, "
    "tương tác với bot thông qua streaming SSE."
)
bul(
    "Zalo User: Người dùng cuối tương tác qua Zalo Bot Direct hoặc Zalo OA, "
    "không cần tài khoản OmniRAG."
)
bul(
    "API Consumer: Tích hợp viên kết nối OmniRAG qua REST API với API key, "
    "sử dụng endpoint POST /api/v1/openrouter/rag/chat hoặc POST /api/v1/bots/{bot_id}/chat."
)
placeholder(
    "Hình 3.2 — Use Case Diagram: Sơ đồ Use Case UML với 4 actor và các use case "
    "chính (Tạo bot, Upload tài liệu, Chat streaming, Tích hợp Zalo, Xem analytics)"
)
ins(
    "Sequence Diagram cho luồng RAG chat streaming mô tả thứ tự tương tác giữa các "
    "thành phần trong khoảng thời gian từ khi nhận request đến khi hoàn tất phản hồi:"
)
bul("t=0ms: Browser gửi POST /api/v1/bots/{bot_id}/chat-stream đến Go Gateway")
bul("t=2ms: Gateway kiểm tra Redis cache — cache miss → proxy đến FastAPI Backend")
bul(
    "t=3ms: Backend song song hóa embed(query) và rewrite(query) qua asyncio.ensure_future()"
)
bul(
    "t=~400ms: Embedding hoàn tất (text-embedding-3-small, 1536 dim) → "
    "khởi động _hybrid_search() ngay lập tức, đồng thời khởi động LightRAG query (nếu domain có KG)"
)
bul(
    "t=~1.300ms: _hybrid_search() trả về top-K*2 chunks sau RRF merge và Cross-Encoder reranking"
)
bul("t=~1.500ms: _rewrite_query() hoàn tất → _crag_classify() bắt đầu với top-3 chunks")
bul("t=~1.700ms: CRAG verdict (relevant/ambiguous/no_context) → lắp ráp context")
bul("t=~2.800ms: LightRAG timeout hoặc hoàn tất → context đầy đủ")
bul("t=~3.500ms: Bắt đầu streaming LLM response qua SSE → browser nhận token đầu tiên")
placeholder(
    "Hình 3.3 — RAG Sequence Diagram: Sơ đồ sequence UML chi tiết các bước từ "
    "request đến streaming response với timeline annotated"
)

# 3.2.2
ins("3.2.2. Biểu đồ hoạt động (Activity Diagram) của pipeline xử lý tài liệu", "Heading 3")
ins(
    "Pipeline xử lý tài liệu (Document Processing Pipeline) diễn ra hoàn toàn bất đồng bộ "
    "qua Celery task queue. Toàn bộ quá trình được chia thành hai giai đoạn độc lập "
    "để đảm bảo bot có thể phục vụ câu hỏi ngay sau khi hoàn thành giai đoạn đầu:"
)
ins(
    "Giai đoạn 1 — Vector Indexing (bắt buộc, ~2-5 phút/tài liệu):"
)
bul("Upload: Frontend POST multipart/form-data → Backend validate (max 25MB, ≤10 định dạng) → lưu MinIO, tạo Document record (status=pending)")
bul("Celery dispatch: process_document_task(document_id, bot_id, file_path, filename)")
bul("Parsing: Download từ MinIO → opendataloader-pdf (format markdown+JSON, hybrid auto/full khi cấu hình bật OCR/picture/formula) → structured chunks kèm page/bbox metadata; fallback sang markdown/local loader nếu JSON thiếu hoặc parser lỗi")
bul("Chunking: Áp dụng chiến lược từ Domain Profile (recursive/sentence/article/parent-child), chunk_size và overlap theo domain")
bul("Contextual prefix: _generate_contextual_prefix_batch() — gọi INTERNAL_LLM_MODEL với 8 chunk/batch (asyncio.gather), sinh 1-2 câu mô tả vị trí chunk trong tài liệu gốc")
bul("Embedding: embed_batch_async() — text đã enriched (prefix + chunk), batch_size=100, gọi OpenRouter text-embedding-3-small (1536-dim)")
bul("Qdrant upsert: Lưu dense OpenRouter vector + sparse BM25 vector + payload {bot_id, document_id, text, parent_text, context_prefix, source, page_numbers, bboxes, element_types, artifact_paths, metadata}")
bul("Cập nhật trạng thái: document.status = 'completed' — bot sẵn sàng phục vụ câu hỏi")
ins(
    "Giai đoạn 2 — Knowledge Graph Indexing (tùy chọn, ~5-30 phút/tài liệu, chạy sau giai đoạn 1):"
)
bul("Điều kiện: enable_knowledge_graph=True (tự động với domain education và legal)")
bul("build_knowledge_graph_task(bot_id, full_text): Sanitize text (xóa null bytes, cắt dòng >10K ký tự, giới hạn 500K ký tự)")
bul("LightRAG.insert_text(): Entity extraction (NER) và Relationship extraction bằng INTERNAL_LLM_MODEL, lưu vào Qdrant + GraphML file")
bul("Cập nhật: doc_metadata.kg_status = 'completed'. Bot KHÔNG bị block trong quá trình này — giai đoạn 1 đã đủ cho chat")
placeholder(
    "Hình 3.4 — Activity Diagram: Sơ đồ hoạt động UML của Document Processing Pipeline "
    "2 giai đoạn (Vector Indexing song song Knowledge Graph Indexing)"
)

# 3.2.3
ins("3.2.3. Sơ đồ luồng tích hợp Zalo Bot Direct và func.vn Hub", "Heading 3")
ins(
    "Hai kênh tích hợp Zalo sử dụng kiến trúc webhook bất đối xứng: Zalo Bot Direct "
    "có webhook riêng theo từng bot ({bot_id} trong URL), trong khi func.vn Hub sử dụng "
    "một webhook tập trung duy nhất cho toàn bộ hệ thống:"
)
ins(
    "Luồng Zalo Bot Direct (token-based, per-bot webhook):"
)
bul("Admin: POST /channels/zalo-bot/connect → hệ thống tự động gọi getMe + setWebhook")
bul("Zalo Platform: POST /channels/zalo-bot/webhook/{bot_id} với x-bot-api-secret-token header")
bul("OmniRAG: HMAC verify → dispatch Celery task → typing indicator → RAG → sendMessage")
ins(
    "Luồng func.vn Hub (centralized webhook, account_id routing):"
)
bul("Admin: Cấu hình OA trong func.vn, trỏ webhook tới POST /channels/zalo/hub-webhook")
bul("Zalo OA: Tin nhắn → func.vn → POST hub-webhook với {account_id, sender_id, content}")
bul("OmniRAG: HMAC verify (x-hub-secret) → lookup bot by account_id → RAG → func.vn Reply API")
ins(
    "Điểm khác biệt kỹ thuật then chốt: Zalo Bot Direct xác thực dựa trên webhook_secret "
    "riêng của từng bot (sinh ngẫu nhiên lúc connect), trong khi func.vn Hub xác thực "
    "bằng một ZALO_HUB_WEBHOOK_SECRET chung toàn hệ thống (cấu hình qua biến môi trường). "
    "Cả hai đều dùng hmac.compare_digest() để ngăn timing attack."
)
placeholder(
    "Hình 3.5 — Zalo Integration Flow: Sơ đồ so sánh luồng Zalo Bot Direct (trái) và "
    "func.vn Hub (phải), highlight sự khác biệt về webhook routing và authentication"
)

# 3.2.4
ins("3.2.4. Biểu đồ luồng xử lý tài liệu 2 giai đoạn: Vector Index và Knowledge Graph", "Heading 3")
ins(
    "Kiến trúc xử lý tài liệu 2 giai đoạn của OmniRAG giải quyết mâu thuẫn giữa "
    "availability (bot sẵn sàng phục vụ sớm) và completeness (KG cần thêm thời gian). "
    "Hai giai đoạn chạy theo mô hình pipeline bất đồng bộ:"
)
bul(
    "Giai đoạn 1 (Critical Path): Upload → MinIO → Celery → Parse → Chunk → Contextual Prefix → "
    "Embed → Qdrant. Trạng thái document chuyển sang 'completed' sau bước này. "
    "Thời gian trung bình: 2-5 phút cho PDF 10MB."
)
bul(
    "Giai đoạn 2 (Background, fire-and-forget): build_knowledge_graph_task được gọi bằng "
    "Celery chain sau khi giai đoạn 1 hoàn thành. LightRAG insert_text() chạy song song "
    "với các request chat đang hoạt động, không block API. "
    "Thời gian: 5-30 phút tùy kích thước tài liệu và tốc độ LLM."
)
bul(
    "Graceful degradation: Nếu giai đoạn 2 thất bại, doc_metadata.kg_status='failed' nhưng "
    "bot vẫn hoạt động bình thường với vector search. KG chỉ là bổ sung, không phải điều kiện bắt buộc."
)
placeholder(
    "Hình 3.6 — 2-Stage Document Processing Flow: Biểu đồ Gantt hoặc pipeline diagram "
    "minh họa thứ tự và tính song song của 2 giai đoạn xử lý"
)

# ── 3.3 ────────────────────────────────────────────────────────────────────
ins("3.3. Thiết kế kiến trúc phần mềm và cơ sở dữ liệu", "Heading 2")
ins(
    "Phần này trình bày thiết kế kỹ thuật chi tiết của hệ thống: kiến trúc phần mềm "
    "đa tầng, lược đồ cơ sở dữ liệu quan hệ, cấu trúc vector collection và lớp "
    "Connector giao tiếp đa kênh."
)

# 3.3.1
ins("3.3.1. Thiết kế kiến trúc Cross-platform với API-first approach", "Heading 3")
ins(
    "OmniRAG áp dụng kiến trúc microservices ba tầng theo nguyên tắc API-first — "
    "mọi giao tiếp giữa các thành phần đều qua REST API có tài liệu OpenAPI đầy đủ:"
)
ins(
    "Tầng 1 — Go API Gateway (cổng 8080): Thành phần đầu tiên nhận mọi request từ "
    "browser, mobile app và Zalo webhook. Được viết bằng Go 1.21 + Gin framework, "
    "Gateway thực hiện ba nhiệm vụ chính không ảnh hưởng đến business logic:"
)
bul(
    "CORS handling: Kiểm soát danh sách origin được phép (frontend localhost:5173, "
    "domain production) theo cấu hình BACKEND_CORS_ORIGINS."
)
bul(
    "Rate limiting: Giới hạn 100 request/giây/IP (cấu hình được qua RATE_LIMIT_RPS). "
    "Dùng Redis Lua script (atomic INCR+EXPIRE) với sliding window 1 giây, "
    "fail-open khi Redis không khả dụng để tránh nghẽn cổ chai."
)
bul(
    "Response caching: Cache GET responses vào Redis (TTL=3600 giây), "
    "key = SHA-256(path|body|authorization) để user A không thấy cache của user B. "
    "Streaming endpoints, analytics và session endpoints bị loại khỏi cache."
)
ins(
    "Tầng 2 — FastAPI Backend (cổng 8000): Xử lý toàn bộ business logic, "
    "được viết bằng Python 3.11 với FastAPI + SQLAlchemy 2.0 async + Pydantic v2. "
    "Kiến trúc dependency injection của FastAPI cho phép inject database session, "
    "current user và các service vào từng endpoint một cách sạch sẽ."
)
bul(
    "Tổ chức router: /api/v1/auth, /bots, /bot-templates, /folders, /analytics, "
    "/dashboard, /users, /integrations, /openrouter, /channels/zalo, /channels/zalo-bot. "
    "Mỗi module endpoint có file riêng trong api/v1/endpoints/."
)
bul(
    "Async throughout: Mọi database call dùng AsyncSession (sqlalchemy.ext.asyncio), "
    "mọi service method đều async để tận dụng tối đa asyncio event loop, "
    "tránh blocking I/O."
)
ins(
    "Tầng 3 — Data Stores (tất cả chạy trong Docker): Bảy dịch vụ lưu trữ phục vụ "
    "các mục đích khác nhau:"
)
bul("PostgreSQL 15 (cổng 5433): Dữ liệu quan hệ — tenants, users, bots, documents, folders")
bul("MongoDB 7.0 (cổng 27017): Dữ liệu phi cấu trúc — conversations, sessions, api_keys, feedback")
bul("Redis 7 (cổng 6380): Celery broker, gateway cache, response cache")
bul("Qdrant (cổng 6333): Vector embeddings + BM25 full-text search")
bul("MinIO (cổng 9000/9001): S3-compatible object storage cho file tài liệu")
bul("OpenDataLoader Hybrid (cổng 5002): PDF parsing với OCR, SmolVLM, LaTeX")
ins(
    "Celery Worker chạy song song với Backend (cùng Docker image, lệnh khác): "
    "xử lý document processing tasks và webhook tasks bất đồng bộ với concurrency=4 workers."
)
placeholder(
    "Hình 3.7 — System Architecture: Sơ đồ kiến trúc tổng thể 3 tầng với Docker services, "
    "ports, và luồng request (Browser → Gateway → Backend → Data Stores)"
)

# 3.3.2
ins("3.3.2. Thiết kế ERD cho PostgreSQL (Users, Tenants, Bots, Documents, Folders)", "Heading 3")
ins(
    "Cơ sở dữ liệu quan hệ PostgreSQL 15 của OmniRAG được thiết kế với năm thực thể "
    "chính, tuân thủ nguyên tắc UUID primary key (bảo mật, distributed-friendly) "
    "và JSONB cho các trường cấu hình linh hoạt (không cần migration khi thêm tham số):"
)
ins(
    "Bảng tenants: Lưu thông tin tổ chức/khách hàng sử dụng OmniRAG. "
    "Các trường: id (UUID, PK), name (VARCHAR 255), email (VARCHAR 255, UNIQUE, INDEX), "
    "plan (VARCHAR 50, DEFAULT='free'), settings (JSONB, DEFAULT={}), "
    "created_at và updated_at (TIMESTAMPTZ). "
    "Quan hệ: 1 tenant có nhiều users và nhiều bots (cascade delete orphans)."
)
ins(
    "Bảng users: Lưu tài khoản người dùng thuộc một tenant. "
    "Các trường: id (UUID, PK), tenant_id (UUID, FK→tenants.id, INDEX), "
    "email (VARCHAR 255, UNIQUE, INDEX), hashed_password (bcrypt), "
    "full_name (VARCHAR 255, nullable), role (VARCHAR 50, DEFAULT='member' — owner/admin/member), "
    "is_active (BOOLEAN, DEFAULT=True), created_at, updated_at."
)
ins(
    "Bảng bots: Lưu cấu hình trợ lý ảo RAG. "
    "Các trường: id (UUID, PK), tenant_id (UUID, FK→tenants.id, INDEX), "
    "name (VARCHAR 255, INDEX), description (TEXT), avatar_url (VARCHAR 500), "
    "config (JSONB — chứa model, system_prompt, temperature, domain, chunking params, "
    "zalo_bot credentials...), api_key (VARCHAR 64, UNIQUE, INDEX), "
    "is_active (BOOLEAN), created_at, updated_at. "
    "Index tổng hợp: (tenant_id, created_at). "
    "Quan hệ: 1 bot có nhiều documents và nhiều folders."
)
ins(
    "Bảng documents: Lưu metadata tài liệu tải lên (file thực tế ở MinIO). "
    "Các trường: id (UUID, PK), bot_id (UUID, FK→bots.id CASCADE, INDEX), "
    "filename (VARCHAR 255), file_path (VARCHAR 500 — MinIO object key), "
    "file_type (VARCHAR 50 — extension normalized), file_size (BIGINT bytes), "
    "content_hash (VARCHAR 64 — SHA256 deduplication), "
    "status (VARCHAR 20 — pending/processing/completed/failed, INDEX), "
    "error_message (TEXT), folder_id (UUID, FK→folders.id SET NULL, nullable), "
    "tags (JSONB list), doc_metadata (JSONB — kg_status, page count...), "
    "created_at (INDEX), updated_at. "
    "Index tổng hợp: (bot_id, created_at), (bot_id, filename), (bot_id, status)."
)
ins(
    "Bảng folders: Hỗ trợ tổ chức tài liệu theo cấu trúc phân cấp (hierarchical). "
    "Các trường: id (UUID, PK), name (VARCHAR 255), "
    "parent_id (UUID, FK→folders.id CASCADE, self-referential, nullable), "
    "bot_id (UUID, FK→bots.id CASCADE, INDEX), created_at, updated_at. "
    "Cấu trúc tự tham chiếu cho phép tạo cây thư mục lồng nhau tùy độ sâu."
)
placeholder(
    "Hình 3.8 — ERD PostgreSQL: Lược đồ quan hệ thực thể đầy đủ với 5 bảng, "
    "khóa chính/ngoại, kiểu dữ liệu và quan hệ cardinality"
)

# 3.3.3
ins("3.3.3. Thiết kế cấu trúc Vector Collection trong Qdrant", "Heading 3")
ins(
    "Qdrant được cấu hình với một collection duy nhất (omnirag) thay vì collection "
    "riêng cho từng bot. Thiết kế này cho phép tái sử dụng HNSW index graph hiệu quả "
    "hơn và dễ quản lý infrastructure. Sự phân tách dữ liệu giữa các bot được thực "
    "hiện hoàn toàn qua payload filtering:"
)
ins(
    "Vector configuration: Mỗi vector point lưu trữ embedding 1536 chiều "
    "(text-embedding-3-small của OpenAI qua OpenRouter), sử dụng Cosine similarity "
    "làm metric khoảng cách, HNSW indexing với m=16 và ef_construct=100 "
    "(cân bằng giữa tốc độ và độ chính xác)."
)
ins(
    "Payload schema của mỗi vector point:"
)
bul("bot_id (string): ID của bot sở hữu chunk — dùng để filter trong mọi query")
bul("text (string): Nội dung chunk đã enriched (context_prefix + original text) — dùng cho BM25 full-text search")
bul("parent_text (string, nullable): Nội dung parent chunk trong Parent-Child strategy — trả về LLM để có ngữ cảnh phong phú hơn")
bul("context_prefix (string): Câu tiêu đề ngữ cảnh 1-2 câu do Contextual Retrieval sinh — lưu riêng để debug")
bul("source (string): Tên file tài liệu gốc — hiển thị trong citation [[n]] cho người dùng")
bul("metadata (object): Thông tin bổ sung: page_number, section, document_id, chunk_index")
ins(
    "Ba payload index tối ưu hóa hiệu năng truy vấn:"
)
bul("KEYWORD index trên bot_id: Lọc nhanh O(1) theo bot — mandatory filter cho mọi vector search")
bul("TEXT index trên text (multilingual tokenizer): Bật BM25 full-text search cho Sparse Retrieval trong Hybrid Search")
bul("KEYWORD index trên source: Cho phép lọc theo tài liệu cụ thể — dùng trong debug retrieval endpoint")
ins(
    "Collection omnirag_memories (Mem0): Collection riêng biệt lưu trữ memory của người dùng. "
    "Metadata: {user_id, bot_id, session_id}. Được quản lý hoàn toàn bởi thư viện mem0ai, "
    "sử dụng cùng vector dimension (1536) và cosine similarity."
)
placeholder(
    "Hình 3.9 — Qdrant Collection Schema: Sơ đồ cấu trúc vector point với payload fields, "
    "3 payload indexes, và luồng upsert từ Document Processing"
)

# 3.3.4
ins("3.3.4. Thiết kế lớp Connector giao tiếp đa kênh (Omnichannel Connector)", "Heading 3")
ins(
    "Lớp Connector trong OmniRAG tách biệt hoàn toàn logic xử lý AI (RAG pipeline) "
    "khỏi logic giao tiếp kênh (channel-specific API calls). Mỗi kênh tích hợp "
    "được đóng gói thành một Service class độc lập với interface thống nhất:"
)
ins(
    "ZaloBotService (backend/app/services/channels/zalo_bot_service.py): "
    "Wrapper cho Zalo Bot Platform API (bot-api.zapps.me). Các method chính:"
)
bul("connect(bot_id, bot_token, webhook_base_url): Toàn bộ flow kết nối — getMe, sinh secret, setWebhook, lưu config")
bul("handle_webhook(bot_id, payload): Xử lý tin nhắn đến — verify bot active, gọi RAG, trả lời")
bul("send_message(bot_token, chat_id, text): Gửi tin nhắn qua POST {TOKEN}/sendMessage")
bul("send_chat_action(bot_token, chat_id, 'typing'): Hiệu ứng đang gõ — cải thiện UX khi RAG xử lý ~3.5s")
bul("_zalo_post(bot_token, method, payload): Wrapper HTTP thống nhất với timeout 15s và error handling")
ins(
    "ZaloHubService (backend/app/services/channels/zalo_hub_service.py): "
    "Connector cho func.vn middleware. Các method chính:"
)
bul("handle_hub_webhook(payload): Nhận payload từ func.vn, tra cứu bot theo account_id, gọi RAG pipeline")
bul("send_reply(user_id, text, oa_id): Gửi phản hồi về func.vn Reply API với Bearer token, timeout 30s")
ins(
    "Cả hai Service đều inject RAG service qua dependency injection — không có coupling "
    "trực tiếp với OpenRouter hay Qdrant. Thiết kế này cho phép dễ dàng thêm kênh mới "
    "(Telegram, Slack, WhatsApp) bằng cách tạo service mới implement cùng interface, "
    "không cần thay đổi RAG pipeline."
)

# ── 3.4 ────────────────────────────────────────────────────────────────────
ins("3.4. Thiết kế giao diện (UI/UX) Dashboard và Chat Playground", "Heading 2")
ins(
    "Frontend OmniRAG xây dựng trên React 19 + TypeScript 5.9, bundler rolldown-vite, "
    "Tailwind CSS 4, Zustand (state management) và TanStack Query (server state). "
    "Tất cả API calls đều qua Axios client có JWT interceptor tự động đính token."
)

# 3.4.1
ins("3.4.1. Dashboard quản trị Tenant và Bot Builder", "Heading 3")
ins(
    "DashboardPage là trang tổng quan dành cho Admin, hiển thị các chỉ số hoạt động "
    "của toàn tenant. Thiết kế ưu tiên parallel data fetching để giảm thời gian load:"
)
bul("4 API calls song song (Promise.all với .catch() graceful degradation): botsApi.list(), analyticsApi.getConversations(8), analyticsApi.getStats(), documentsApi.list(botId)")
bul("Stats tiles: Tổng số bot, phiên chat hoạt động, tin nhắn hôm nay, thời gian phản hồi trung bình")
bul("Danh sách hội thoại gần đây: 8 cuộc hội thoại mới nhất với bot name, tin nhắn tóm tắt, timestamp")
bul("Agent status: Hiển thị trạng thái các bot đang hoạt động và Celery worker status")
ins(
    "BotWizardPage triển khai multi-step wizard với bốn bước có progress indicator rõ ràng. "
    "Domain Selector (bước 2) hiển thị 4 cards với màu sắc đặc trưng (getDomainMeta() "
    "từ utils/domainHelpers.ts), mô tả use case và ví dụ tài liệu phù hợp cho mỗi domain. "
    "BotConfigPage luôn hiển thị toàn bộ advanced controls (không có Simple/Advanced toggle) "
    "để người dùng kiểm soát đầy đủ ngay từ đầu. Sau khi upload tài liệu thành công, "
    "CTA 'Start Chatting' xuất hiện để dẫn người dùng vào Chat Playground."
)
placeholder(
    "Hình 3.10 — Dashboard và Bot Wizard: Screenshot giao diện Dashboard (trái) "
    "và màn hình Domain Selector trong Bot Wizard (phải)"
)

# 3.4.2
ins("3.4.2. Giao diện cấu hình kênh tích hợp (Channel Configuration)", "Heading 3")
ins(
    "Giao diện tích hợp kênh được thiết kế tối giản hóa quá trình kết nối — "
    "toàn bộ chỉ cần một bước paste Bot Token:"
)
bul("Zalo Bot Direct: Form nhập Bot Token → nhấn Connect → hệ thống tự gọi getMe (verify) + setWebhook → hiển thị trạng thái Connected, Webhook URL và thời điểm kết nối")
bul("Zalo Hub: Form nhập Zalo OA Account ID (account_id từ func.vn admin) → Save → OmniRAG lưu vào bot config để routing webhook tập trung")
bul("Disconnect: Nút ngắt kết nối → xóa credentials khỏi bot config, bot chuyển sang trạng thái Disconnected")
ins(
    "Trạng thái kết nối được lấy real-time qua GET /api/v1/channels/zalo-bot/status/{bot_id}, "
    "hiển thị badge màu xanh/đỏ tương ứng. Webhook URL được hiển thị để người dùng có thể "
    "copy và kiểm tra thủ công nếu cần debug."
)
placeholder(
    "Hình 3.11 — Channel Configuration UI: Screenshot màn hình tích hợp Zalo Bot Direct "
    "với trạng thái Connected và Webhook URL display"
)

# 3.4.3
ins("3.4.3. Knowledge Graph Visualization: sigma.js và graphology trên React", "Heading 3")
ins(
    "KnowledgeGraphPage cung cấp giao diện trực quan hóa đồ thị tri thức tương tác, "
    "giúp người dùng hiểu cấu trúc knowledge base của bot thay vì chỉ xem danh sách tài liệu. "
    "Dữ liệu đồ thị được lấy từ GET /api/v1/bots/{bot_id}/knowledge-graph, "
    "trả về định dạng {nodes, links} từ LightRAG GraphML."
)
bul("Thư viện: @react-sigma/core (wrapper React cho sigma.js v3) + graphology (cấu trúc dữ liệu đồ thị)")
bul("Nodes: Entity được trích xuất từ tài liệu (tên, khái niệm, địa điểm...) — kích thước node tỷ lệ với số lần xuất hiện")
bul("Edges: Quan hệ giữa entities (has-property, related-to, causes...) — màu/độ dày edge theo weight")
bul("Layout: Force-directed algorithm (FA2) — entity liên quan nhiều sẽ tự động tập hợp gần nhau")
bul("Interaction: Hover hiển thị description, click highlight subgraph kết nối, zoom/pan tự do")
ins(
    "Visualization trực quan giúp người dùng phát hiện entity extraction chất lượng thấp "
    "(entity bị trích xuất sai, quan hệ không chính xác), từ đó điều chỉnh tài liệu hoặc "
    "cấu hình LightRAG extraction model phù hợp hơn."
)
placeholder(
    "Hình 3.12 — Knowledge Graph Visualization: Screenshot KnowledgeGraphPage với force-directed "
    "graph, node labels và relationship edges"
)

# 3.4.4
ins("3.4.4. Streaming Chat UI: SSE, Session management và Message Feedback", "Heading 3")
ins(
    "ChatPage là giao diện tương tác chính của người dùng cuối, thiết kế tập trung "
    "vào trải nghiệm real-time và transparency về nguồn thông tin:"
)
bul(
    "SSE Streaming: Kết nối đến POST /api/v1/bots/{bot_id}/chat-stream qua EventSource API. "
    "Mỗi token LLM được stream ngay khi sinh, tạo cảm giác phản hồi nhanh dù "
    "total latency là 3.5-4.5 giây. Spinner hiển thị trong giai đoạn 'Thinking' "
    "trước khi token đầu tiên đến."
)
bul(
    "Citation rendering: Câu trả lời có thể chứa [[1]], [[2]]... — frontend render thành "
    "badge có thể click, hiển thị tên file nguồn (source) và đoạn text tương ứng. "
    "Người dùng có thể kiểm tra độ tin cậy của thông tin."
)
bul(
    "Session management: Mỗi cuộc hội thoại có session_id (UUID). "
    "Sidebar hiển thị danh sách sessions với tiêu đề tự động sinh bởi LLM "
    "từ tin nhắn đầu tiên. Người dùng có thể switch giữa các session và xóa session."
)
bul(
    "Message feedback: Mỗi câu trả lời có nút thumbs up/down. "
    "Feedback được lưu vào MongoDB message_feedback collection để phân tích chất lượng RAG "
    "và hỗ trợ RLHF (Reinforcement Learning from Human Feedback) trong tương lai."
)
bul(
    "Agent logs: Nút 'Show reasoning' mở panel hiển thị từng bước pipeline "
    "(Analyzing Query → Knowledge Retrieval → CRAG Check → Generating Answer) "
    "với timestamp, tăng tính minh bạch cho người dùng kỹ thuật."
)
placeholder(
    "Hình 3.13 — Chat UI: Screenshot ChatPage với streaming response, citation badges, "
    "session sidebar và agent log panel"
)

# ── Lưu file ────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"\nChương 3 đã chèn thành công → {OUT}")
print(f"Backup gốc: {BAK}")
print(f"\nVui lòng mở {OUT} bằng Word/LibreOffice để kiểm tra trước khi chạy pass 2.")
