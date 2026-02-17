
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def set_page_border(section, color="2E598A", val="double"):
    sec_pr = section._sectPr
    for pb in sec_pr.findall(qn('w:pgBorders')):
        sec_pr.remove(pb)
    pg_borders = OxmlElement('w:pgBorders')
    pg_borders.set(qn('w:offsetFrom'), 'page')
    for border_name in ['top', 'left', 'bottom', 'right']:
        edge = OxmlElement(f'w:{border_name}')
        edge.set(qn('w:val'), val) 
        edge.set(qn('w:sz'), '12')      
        edge.set(qn('w:space'), '24')
        edge.set(qn('w:color'), color)
        pg_borders.append(edge)
    sec_pr.append(pg_borders)

def remove_page_border(section):
    sec_pr = section._sectPr
    for pb in sec_pr.findall(qn('w:pgBorders')):
        sec_pr.remove(pb)
    pg_borders = OxmlElement('w:pgBorders')
    pg_borders.set(qn('w:offsetFrom'), 'page')
    for border_name in ['top', 'left', 'bottom', 'right']:
        edge = OxmlElement(f'w:{border_name}')
        edge.set(qn('w:val'), 'none')
        edge.set(qn('w:sz'), '0')
        edge.set(qn('w:space'), '0')
        pg_borders.append(edge)
    sec_pr.append(pg_borders)

def add_toc_entry(doc, text, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3 * level)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    if level == 0: run.bold = True
    p.add_run('\t') 
    p.add_run('...') 

def create_de_cuong_final():
    doc = Document()
    FINAL_TITLE = "XÂY DỰNG NỀN TẢNG OMNIRAG: GIẢI PHÁP ĐA NỀN TẢNG HỖ TRỢ TỰ KHỞI TẠO TRỢ LÝ ẢO THÔNG MINH ĐA KÊNH TỪ DỮ LIỆU TÙY CHỈNH CHO CÁ NHÂN VÀ TỔ CHỨC"
    
    # --- COVER PAGES ---
    section = doc.sections[0]
    section.top_margin = Inches(0.79)
    section.bottom_margin = Inches(0.79)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(0.79)
    set_page_border(section, color="2E598A") 

    # PAGE 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BỘ GIÁO DỤC VÀ ĐÀO TẠO\nTRƯỜNG ĐẠI HỌC MỎ - ĐỊA CHẤT\n---------------------------------------')
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph('\n')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[HỌ VÀ TÊN SINH VIÊN]')
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph('\n' * 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ĐỒ ÁN TỐT NGHIỆP\nNGÀNH CÔNG NGHỆ THÔNG TIN')
    run.bold = True
    run.font.size = Pt(18)
    doc.add_paragraph('\n')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TÊN ĐỀ TÀI:\n')
    run.bold = True
    run.font.size = Pt(14)
    run = p.add_run(FINAL_TITLE)
    run.bold = True
    run.font.size = Pt(17)
    doc.add_paragraph('\n' * 3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Hà Nội - Năm 2026')
    run.font.size = Pt(12)
    doc.add_page_break()

    # PAGE 2
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BỘ GIÁO DỤC VÀ ĐÀO TẠO\nTRƯỜNG ĐẠI HỌC MỎ - ĐỊA CHẤT\n---------------------------------------')
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph('\n' * 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ĐỒ ÁN TỐT NGHIỆP\nCHUYÊN NGÀNH: KHOA HỌC MÁY TÍNH')
    run.bold = True
    run.font.size = Pt(18)
    doc.add_paragraph('\n' * 3)
    table = doc.add_table(rows=4, cols=2)
    table.cell(0,0).text = "GIÁO VIÊN HƯỚNG DẪN:"
    table.cell(0,1).text = "ThS. [Tên Thầy/Cô]"
    table.cell(1,0).text = "SINH VIÊN THỰC HIỆN:"
    table.cell(1,1).text = "[Tên của Bro]"
    table.cell(2,0).text = "BỘ MÔN:"
    table.cell(2,1).text = "KHOA HỌC MÁY TÍNH"
    table.cell(3,0).text = "LỚP:"
    table.cell(3,1).text = "KHMT ỨNG DỤNG [Lớp]"
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(13)
                    r.bold = True
    doc.add_paragraph('\n' * 8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Hà Nội – Năm 2026')
    run.font.size = Pt(12)

    # --- TABLE OF CONTENTS (FULL VERSION) ---
    new_sec = doc.add_section()
    remove_page_border(new_sec)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('MỤC LỤC')
    run.bold = True
    run.font.size = Pt(16)
    
    add_toc_entry(doc, "LỜI CẢM ƠN", 0)
    add_toc_entry(doc, "DANH MỤC KÝ HIỆU VÀ CHỮ VIẾT TẮT", 0)
    add_toc_entry(doc, "DANH MỤC BẢNG BIỂU", 0)
    add_toc_entry(doc, "DANH MỤC HÌNH ẢNH", 0)
    add_toc_entry(doc, "MỞ ĐẦU", 0)
    add_toc_entry(doc, "1. Lý do chọn đề tài", 1)
    add_toc_entry(doc, "2. Mục tiêu của đề tài", 1)
    add_toc_entry(doc, "3. Đối tượng và phạm vi nghiên cứu", 1)
    add_toc_entry(doc, "4. Phương pháp nghiên cứu", 1)
    add_toc_entry(doc, "5. Bố cục đề tài", 1)
    
    # CHƯƠNG 1 - FULL VERSION
    add_toc_entry(doc, "CHƯƠNG 1. TỔNG QUAN VÀ KHẢO SÁT HỆ THỐNG", 0)
    add_toc_entry(doc, "1.1. Giới thiệu bài toán trợ lý ảo thông minh cho cá nhân và tổ chức", 1)
    add_toc_entry(doc, "1.2. Nhu cầu giao tiếp đa kênh (Omnichannel) trong doanh nghiệp hiện đại", 1)
    add_toc_entry(doc, "1.3. Thách thức trong triển khai trợ lý ảo đa nền tảng và đa nguồn dữ liệu", 1)
    add_toc_entry(doc, "1.4. Khảo sát các giải pháp hiện nay", 1)
    add_toc_entry(doc, "1.4.1. Các nền tảng chatbot đóng gói sẵn (No-code/Low-code Platforms)", 2)
    add_toc_entry(doc, "1.4.2. Các hạn chế về tính riêng tư và tùy biến dữ liệu", 2)
    add_toc_entry(doc, "1.4.3. Đánh giá khả năng tích hợp đa kênh của các giải pháp hiện tại", 2)
    add_toc_entry(doc, "1.5. Phân tích yêu cầu hệ thống", 1)
    add_toc_entry(doc, "1.5.1. Yêu cầu chức năng: Quản lý Tenant, Bot Builder, Multi-channel Chat", 2)
    add_toc_entry(doc, "1.5.2. Yêu cầu phi chức năng: Bảo mật, Hiệu năng, Khả năng mở rộng", 2)
    add_toc_entry(doc, "1.5.3. Yêu cầu về khả năng tích hợp bên thứ ba (Third-party Integration)", 2)
    
    # CHƯƠNG 2 - FULL VERSION
    add_toc_entry(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VỀ AI VÀ RAG NÂNG CAO", 0)
    add_toc_entry(doc, "2.1. Tổng quan về Large Language Models (LLM) và Generative AI", 1)
    add_toc_entry(doc, "2.2. Kiến trúc Retrieval-Augmented Generation (RAG)", 1)
    add_toc_entry(doc, "2.2.1. Quy trình RAG cơ bản và các điểm nghẽn (Bottlenecks)", 2)
    add_toc_entry(doc, "2.2.2. Phân loại RAG: Naive RAG vs Advanced RAG vs Agentic RAG", 2)
    add_toc_entry(doc, "2.3. Các kỹ thuật tối ưu hóa truy xuất nâng cao", 1)
    add_toc_entry(doc, "2.3.1. Kỹ thuật Hybrid Search (Dense & Sparse Retrieval)", 2)
    add_toc_entry(doc, "2.3.2. Query Transformation: HyDE và Multi-Query Generation", 2)
    add_toc_entry(doc, "2.3.3. Document Re-ranking và áp dụng Cross-Encoders", 2)
    add_toc_entry(doc, "2.3.4. Các chiến lược Chunking tối ưu: Recursive và Semantic Splitting", 2)
    add_toc_entry(doc, "2.4. Khái niệm AI Agents và Agentic Workflow", 1)
    add_toc_entry(doc, "2.4.1. Cơ chế Reasoning và Acting (ReAct Pattern)", 2)
    add_toc_entry(doc, "2.4.2. Quản lý trạng thái và luồng hội thoại phức tạp", 2)
    add_toc_entry(doc, "2.4.3. Multi-Agent Collaboration và Task Delegation", 2)
    
    # CHƯƠNG 3 - FULL VERSION WITH OMNICHANNEL
    add_toc_entry(doc, "CHƯƠNG 3. CÁC CÔNG NGHỆ VÀ FRAMEWORK TRỌNG TÂM", 0)
    add_toc_entry(doc, "3.1. Ecosystem LangChain cho phát triển AI Applications", 1)
    add_toc_entry(doc, "3.1.1. LangChain: Xây dựng chuỗi (Chains) và tích hợp Tools", 2)
    add_toc_entry(doc, "3.1.2. LangGraph: Điều phối Agents phức tạp với Cyclic Graph", 2)
    add_toc_entry(doc, "3.1.3. LangFuse: Tracing, Observability và Evaluation chất lượng RAG", 2)
    add_toc_entry(doc, "3.2. Cơ sở dữ liệu và hạ tầng lưu trữ", 1)
    add_toc_entry(doc, "3.2.1. Vector Database Qdrant và HNSW Indexing", 2)
    add_toc_entry(doc, "3.2.2. SQL (PostgreSQL) và NoSQL (MongoDB) trong hệ thống Hybrid", 2)
    add_toc_entry(doc, "3.2.3. Redis: Caching và Session Management", 2)
    add_toc_entry(doc, "3.2.4. MinIO: Object Storage tương thích S3", 2)
    add_toc_entry(doc, "3.3. Phát triển đa nền tảng và quản trị hệ thống", 1)
    add_toc_entry(doc, "3.3.1. Backend FastAPI: Bất đồng bộ và hiệu năng cao", 2)
    add_toc_entry(doc, "3.3.2. Hạ tầng Docker Microservices và Celery Task Queue", 2)
    add_toc_entry(doc, "3.3.3. Thiết kế Responsive Web với React và Tailwind CSS", 2)
    add_toc_entry(doc, "3.4. Công nghệ tích hợp đa kênh (Omnichannel Integration)", 1)
    add_toc_entry(doc, "3.4.1. func.vn Gateway: Adapter cho Zalo OA và Facebook Messenger", 2)
    add_toc_entry(doc, "3.4.2. Telegram Bot API: Webhook và Long Polling", 2)
    add_toc_entry(doc, "3.4.3. Phát triển Web Widget nhúng thông minh (Embeddable Chat Widget)", 2)
    add_toc_entry(doc, "3.4.4. API Gateway Pattern cho đa nền tảng", 2)
    
    # CHƯƠNG 4 - FULL VERSION
    add_toc_entry(doc, "CHƯƠNG 4. PHÂN TÍCH VÀ THIẾT KẾ NỀN TẢNG OMNIRAG", 0)
    add_toc_entry(doc, "4.1. Phân tích chức năng và luồng nghiệp vụ", 1)
    add_toc_entry(doc, "4.1.1. Quy trình Self-build bot từ PDF, Website và API", 2)
    add_toc_entry(doc, "4.1.2. Cơ chế phân tách dữ liệu đa khách hàng (Isolated Multi-tenancy)", 2)
    add_toc_entry(doc, "4.1.3. Luồng xử lý tin nhắn đa kênh (Omnichannel Message Flow)", 2)
    add_toc_entry(doc, "4.2. Thiết kế các biểu đồ hệ thống", 1)
    add_toc_entry(doc, "4.2.1. Biểu đồ Use Case và Sequence Diagrams cho RAG Agent", 2)
    add_toc_entry(doc, "4.2.2. Biểu đồ hoạt động (Activity Diagram) của pipeline xử lý dữ liệu", 2)
    add_toc_entry(doc, "4.2.3. Sơ đồ luồng tích hợp đa kênh qua func.vn", 2)
    add_toc_entry(doc, "4.3. Thiết kế kiến trúc phần mềm và cơ sở dữ liệu", 1)
    add_toc_entry(doc, "4.3.1. Thiết kế kiến trúc Cross-platform với API-first approach", 2)
    add_toc_entry(doc, "4.3.2. Thiết kế ERD cho PostgreSQL (Users, Tenants, Bots, Documents)", 2)
    add_toc_entry(doc, "4.3.3. Thiết kế cấu trúc Vector Collection trong Qdrant", 2)
    add_toc_entry(doc, "4.3.4. Thiết kế lớp Connector giao tiếp đa kênh (Omnichannel Connector)", 2)
    add_toc_entry(doc, "4.4. Thiết kế giao diện (UI/UX) Dashboard và Chat Playground", 1)
    add_toc_entry(doc, "4.4.1. Dashboard quản trị Tenant và Bot Builder", 2)
    add_toc_entry(doc, "4.4.2. Giao diện cấu hình kênh tích hợp (Channel Configuration)", 2)
    
    # CHƯƠNG 5 - FULL VERSION
    add_toc_entry(doc, "CHƯƠNG 5. CÀI ĐẶT, THỰC NGHIỆM VÀ ĐÁNH GIÁ KẾT QUẢ", 0)
    add_toc_entry(doc, "5.1. Triển khai và cài đặt hệ thống", 1)
    add_toc_entry(doc, "5.1.1. Triển khai hạ tầng Microservices trên Docker", 2)
    add_toc_entry(doc, "5.1.2. Cấu hình LangGraph Agent workflow xử lý hội thoại", 2)
    add_toc_entry(doc, "5.1.3. Tích hợp func.vn cho Zalo OA và Facebook Messenger", 2)
    add_toc_entry(doc, "5.1.4. Cài đặt Telegram Bot và Web Widget", 2)
    add_toc_entry(doc, "5.2. Kết quả thực nghiệm và Demo", 1)
    add_toc_entry(doc, "5.2.1. Demo giao diện Web Dashboard và Bot Chat", 2)
    add_toc_entry(doc, "5.2.2. Thử nghiệm trợ lý ảo trên Zalo OA với dữ liệu tùy chỉnh", 2)
    add_toc_entry(doc, "5.2.3. Thử nghiệm Web Widget nhúng trên đa nền tảng Website", 2)
    add_toc_entry(doc, "5.2.4. Thử nghiệm khả năng tự học từ đa định dạng dữ liệu", 2)
    add_toc_entry(doc, "5.3. Đánh giá và kiểm thử chất lượng", 1)
    add_toc_entry(doc, "5.3.1. Đánh giá tốc độ phản hồi (Latency) với Redis Cache", 2)
    add_toc_entry(doc, "5.3.2. Đánh giá độ chính xác (RAG Metrics) thông qua LangFuse", 2)
    add_toc_entry(doc, "5.3.3. Đánh giá tính nhất quán đa kênh (Cross-channel Consistency)", 2)
    add_toc_entry(doc, "5.4. Thảo luận và phân tích kết quả", 1)
    
    add_toc_entry(doc, "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 0)
    add_toc_entry(doc, "DANH MỤC TÀI LIỆU THAM KHẢO", 0)
    add_toc_entry(doc, "PHỤ LỤC", 0)
    add_toc_entry(doc, "Phụ lục A: Danh mục API hỗ trợ đa nền tảng và Third-party integration", 1)
    add_toc_entry(doc, "Phụ lục B: Quy trình thiết lập hệ thống Monitoring với LangFuse", 1)
    add_toc_entry(doc, "Phụ lục C: Các kịch bản kiểm thử (Test Cases) và kết quả chi tiết", 1)
    add_toc_entry(doc, "Phụ lục D: Hướng dẫn cấu hình func.vn cho Zalo OA và Facebook", 1)
    add_toc_entry(doc, "Phụ lục E: Mã nguồn mẫu tích hợp Web Widget và Telegram Bot", 1)

    doc.add_page_break()

    # --- CONTENT (Sample) ---
    doc.add_heading('MỞ ĐẦU', level=1)
    doc.add_paragraph("Nội dung chi tiết của đề cương OmniRAG - Nền tảng đa kênh hỗ trợ tự khởi tạo trợ lý ảo từ dữ liệu tùy chỉnh...")
    
    doc.save('OmniRAG_DeCuong_FINAL_COMPLETE_V8.docx')
    print("DONE: OmniRAG_DeCuong_FINAL_COMPLETE_V8.docx")

if __name__ == "__main__":
    create_de_cuong_final()
