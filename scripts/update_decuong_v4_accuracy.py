#!/usr/bin/env python3
"""Patch Decuong OmniRAG v4 text and embedded RAG diagrams without changing layout."""

from __future__ import annotations

import io
import shutil
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "docs" / "Decuong_OmniRAG_v4.docx"
OUT = ROOT / "docs" / "Decuong_OmniRAG_v4_fixed.docx"
ASSET_DIR = ROOT / "docs" / "diagrams" / "docx_replacements"


REPLACEMENTS = {
    "Về bảo mật: JWT authentication (HS256, 30 phút), bcrypt password hashing, HMAC constant-time comparison cho webhook verification, UUID-based file storage key (ngăn path traversal), MIME type validation cho file upload, phân quyền theo role (owner/admin/member).":
        "Về bảo mật: JWT authentication (HS256, 30 phút), bcrypt password hashing, constant-time comparison cho webhook secret/token verification, UUID-based object key để ngăn path traversal, kiểm tra extension + content-type sanity-check khi upload file, và phân quyền theo role (owner/admin/member).",
    "Về hiệu năng: Go API Gateway xử lý 40.000+ request/giây cho cached requests (2ms latency); Redis cache với TTL 1 giờ giảm tải LLM API; Celery worker xử lý tài liệu bất đồng bộ; Qdrant Scalar Quantization INT8 giảm memory footprint; batch embedding với ThreadPoolExecutor.":
        "Về hiệu năng: Go API Gateway tách lớp proxy/cache GET khỏi backend; Redis cache giảm tải embedding/rewrite/CRAG/chat stateless; Celery worker xử lý tài liệu bất đồng bộ; Qdrant v1.16 lưu dense OpenRouter vector và sparse BM25 vector trong collection v3; benchmark Docker isolated gần nhất ingest 6 fixture thành công, tạo 6 Qdrant points và trả về 5 retrieved chunks cho truy vấn thử nghiệm.",
    "Trong OmniRAG, Qdrant được cấu hình với Full-Text Search index (MULTILINGUAL tokenizer) cho BM25-style search song song với HNSW vector index. Kết quả từ hai luồng được hợp nhất bằng Reciprocal Rank Fusion (RRF) với hệ số k=60: score(d) = Σ 1/(k + rank_i(d)). Sau đó, Cross-Encoder BAAI/bge-reranker-v2-m3 rerank top candidates, điểm được normalize qua sigmoid để đưa về khoảng [0,1].":
        "Trong OmniRAG, Qdrant được cấu hình theo collection v3 với named vectors: dense vector 1536 chiều từ OpenRouter embedding và sparse BM25 vector sinh bởi FastEmbed với IDF modifier. Query API thực hiện prefetch song song dense+sparse, hợp nhất kết quả bằng Reciprocal Rank Fusion (RRF), sau đó Cross-Encoder rerank top candidates; nếu reranker không khả dụng, hệ thống fallback về RRF score.",
    "Qdrant là vector database mã nguồn mở viết bằng Rust, hỗ trợ HNSW (Hierarchical Navigable Small World) indexing cho approximate nearest neighbor search với độ phức tạp O(log n). OmniRAG cấu hình Qdrant với Scalar Quantization INT8 (giảm memory 4x so với float32) và always_ram=True để giữ quantized vectors trong RAM, đảm bảo latency thấp.":
        "Qdrant là vector database mã nguồn mở viết bằng Rust, hỗ trợ HNSW (Hierarchical Navigable Small World) indexing cho approximate nearest neighbor search. OmniRAG hiện pin Qdrant server v1.16.0 và dùng collection omnirag_openrouter_collection_v3 với hai named vectors: dense cosine 1536 chiều và sparse BM25 có IDF modifier. Cấu hình này ưu tiên hybrid retrieval đúng nghĩa thay vì dựa vào full-text MatchText.",
    "Payload schema của mỗi vector point: {bot_id, source, text, parent_text, context_prefix, metadata}. Có 3 index: KEYWORD index trên bot_id (filter nhanh theo tenant), TEXT index trên text field (Full-Text Search), và composite index bot_id+status/bot_id+created_at trong PostgreSQL (Alembic migration 69b59b9d) cho document listing queries.":
        "Payload schema của mỗi vector point gồm: bot_id, document_id, source, text, parent_text, context_prefix, page_numbers, bboxes, element_types, artifact_paths và metadata. Qdrant tạo payload index cho bot_id, document_id và source để filter nhanh; sparse BM25 được lưu như named sparse vector, không dựa vào full-text payload search.",
    "•  Tầng vector (Qdrant): Tất cả vector embedding được lưu trong một collection duy nhất (omnirag) nhưng mỗi document chunk mang payload {bot_id} làm định danh phân tách. Mọi truy vấn vector search đều đính kèm bộ lọc bot_id, đảm bảo bot A không bao giờ truy xuất được dữ liệu của bot B dù cùng chung hạ tầng.":
        "•  Tầng vector (Qdrant): Tất cả dense/sparse vector được lưu trong collection omnirag_openrouter_collection_v3, nhưng mỗi document chunk mang payload {bot_id, document_id} làm định danh phân tách. Mọi truy vấn retrieval đều đính kèm bộ lọc bot_id, đảm bảo bot A không truy xuất dữ liệu của bot B dù dùng chung hạ tầng.",
    "•  Parsing: Download từ MinIO → opendataloader-pdf (hybrid mode: OCR + SmolVLM image description + LaTeX formula) → markdown text. Fallback: docling-fast → local → PyPDFLoader":
        "•  Parsing: Download từ MinIO → opendataloader-pdf với format markdown+JSON, table_method=cluster, reading_order=xycut, image_output=external → structured chunks kèm page/bbox/element metadata. Nếu JSON thiếu hoặc parser lỗi, backend fallback sang markdown/local loader thay vì reset collection.",
    "•  Tổ chức router: /api/v1/auth, /bots, /bot-templates, /folders, /analytics, /dashboard, /users, /integrations, /openrouter, /channels/zalo, /channels/zalo-bot. Mỗi module endpoint có file riêng trong api/v1/endpoints/.":
        "•  Tổ chức router: /api/v1/auth, /bots, /bot-templates, /folders, /analytics, /dashboard, /users, /integrations, /openrouter, /channels/zalo, /channels/zalo-bot, /channels/zalo-personal, /channels/facebook và /channels/telegram. Mỗi module endpoint có file riêng trong api/v1/endpoints/.",
    "•  Qdrant (cổng 6333): Vector embeddings + BM25 full-text search":
        "•  Qdrant (cổng 6333): Dense embeddings + sparse BM25 named vectors cho RRF hybrid retrieval",
    "•  OpenDataLoader Hybrid (cổng 5002): PDF parsing với OCR, SmolVLM, LaTeX":
        "•  OpenDataLoader Hybrid (cổng 5002): CPU-only PDF layout/OCR parsing, optional picture/formula enrichment",
    "Bảng documents: Lưu metadata tài liệu tải lên (file thực tế ở MinIO). Các trường: id (UUID, PK), bot_id (UUID, FK→bots.id CASCADE, INDEX), filename (VARCHAR 255), file_path (VARCHAR 500 — MinIO object key), file_type (VARCHAR 50 — MIME type), file_size (BIGINT bytes), content_hash (VARCHAR 64 — SHA256 deduplication), status (VARCHAR 20 — pending/processing/completed/failed, INDEX), error_message (TEXT), folder_id (UUID, FK→folders.id SET NULL, nullable), tags (JSONB list), doc_metadata (JSONB — kg_status, page count...), created_at (INDEX), updated_at. Index tổng hợp: (bot_id, created_at), (bot_id, filename), (bot_id, status).":
        "Bảng documents: Lưu metadata tài liệu tải lên (file thực tế ở MinIO). Các trường: id (UUID, PK), bot_id (UUID, FK→bots.id CASCADE, INDEX), filename (VARCHAR 255), file_path (VARCHAR 500 — MinIO object key), file_type (VARCHAR 50 — extension normalized như pdf/docx/xlsx), file_size (BIGINT bytes), content_hash (VARCHAR 64 — SHA256 deduplication), status (VARCHAR 20 — pending/processing/completed/failed, INDEX), error_message (TEXT), folder_id (UUID, FK→folders.id SET NULL, nullable), tags (JSONB list), doc_metadata (JSONB — content_type, kg_status, page/chunk/artifact metadata...), created_at (INDEX), updated_at. Index tổng hợp: (bot_id, created_at), (bot_id, filename), (bot_id, status).",
    "Qdrant được cấu hình với một collection duy nhất (omnirag) thay vì collection riêng cho từng bot. Thiết kế này cho phép tái sử dụng HNSW index graph hiệu quả hơn và dễ quản lý infrastructure. Sự phân tách dữ liệu giữa các bot được thực hiện hoàn toàn qua payload filtering:":
        "Qdrant được cấu hình với một collection mặc định omnirag_openrouter_collection_v3 thay vì collection riêng cho từng bot. Thiết kế này cho phép tái sử dụng index hiệu quả hơn và dễ quản lý infrastructure. Sự phân tách dữ liệu giữa các bot được thực hiện hoàn toàn qua payload filtering:",
    "Vector configuration: Mỗi vector point lưu trữ embedding 1536 chiều (text-embedding-3-small của OpenAI qua OpenRouter), sử dụng Cosine similarity làm metric khoảng cách, HNSW indexing với m=16 và ef_construct=100 (cân bằng giữa tốc độ và độ chính xác).":
        "Vector configuration: Mỗi point lưu dense vector 1536 chiều (text-embedding-3-small qua OpenRouter) với Cosine similarity và sparse BM25 vector sinh bởi FastEmbed với IDF modifier. Dense+sparse prefetch được fuse bằng RRF trong Qdrant Query API.",
    "•  text (string): Nội dung chunk đã enriched (context_prefix + original text) — dùng cho BM25 full-text search":
        "•  text (string): Nội dung chunk đã enriched (context_prefix + original text) — dùng cho display, rerank và LLM context",
    "•  metadata (object): Thông tin bổ sung: page_number, section, document_id, chunk_index":
        "•  metadata (object): Thông tin bổ sung: chunk_index, heading_path, opendataloader_element_ids, has_structured_json và metadata nguồn",
    "Ba payload index tối ưu hóa hiệu năng truy vấn:":
        "Các payload index tối ưu hóa hiệu năng truy vấn:",
    "•  TEXT index trên text (multilingual tokenizer): Bật BM25 full-text search cho Sparse Retrieval trong Hybrid Search":
        "•  KEYWORD index trên document_id: Lọc/debug theo tài liệu và hỗ trợ delete/reindex chính xác",
    "•  KEYWORD index trên source: Cho phép lọc theo tài liệu cụ thể — dùng trong debug retrieval endpoint":
        "•  KEYWORD index trên source: Cho phép lọc theo file nguồn — dùng trong debug retrieval endpoint",
    "•  qdrant (Qdrant, port 6333): Volume qdrant_data. Collection omnirag tự động tạo với HNSW index khi backend upsert lần đầu.":
        "•  qdrant (Qdrant v1.16.0, port 6333): Volume qdrant_data. Collection omnirag_openrouter_collection_v3 tự động tạo khi backend upsert lần đầu, gồm dense vector và sparse BM25 named vector.",
    "•  opendataloader-hybrid (port 5002): Service OCR và image description. Dockerfile: backend/Dockerfile.hybrid. Yêu cầu Java 21+ và mô hình SmolVLM. Cung cấp endpoint /parse-pdf cho backend.":
        "•  opendataloader-hybrid (port 5002): Service CPU-only cho PDF layout/OCR parsing và optional picture/formula enrichment. Dockerfile: backend/Dockerfile.hybrid cài CPU-only torch trước khi cài opendataloader-pdf[hybrid] để tránh CUDA wheels không cần thiết.",
    "•  Bước 2 — Hybrid Search (t≈400ms, khởi động ngay khi embedding xong): _hybrid_search() thực hiện song song Vector Search (cosine similarity, Qdrant HNSW) và Full-Text Search (BM25-style, Qdrant TEXT index, multilingual tokenizer). Kết quả được hợp nhất bằng Reciprocal Rank Fusion (RRF, k=60) rồi rerank bằng Cross-Encoder ms-marco-MiniLM-L-6-v2.":
        "•  Bước 2 — Hybrid Search: _hybrid_search() gọi Qdrant query_points với prefetch dense vector và sparse BM25 vector, fuse bằng Reciprocal Rank Fusion (RRF), sau đó rerank bằng Cross-Encoder ms-marco-MiniLM-L-6-v2. Benchmark isolated gần nhất ghi nhận retrieval khoảng 3.0 giây trên fixture nhỏ, bao gồm embedding/retrieval/rerank.",
    "Kịch bản 2 — Photo with Vision: Người dùng chụp ảnh màn hình lỗi phần mềm và gửi kèm caption 'Làm sao để sửa lỗi này?'. Bot: (a) Download ảnh từ Telegram server (file_id → getFile → download_file); (b) Encode base64, tạo data URL; (c) Gọi OpenRouter vision model (gpt-4o-mini) với prompt 'Describe this image briefly in Vietnamese (2-3 sentences)' + user caption; (d) Feed kết quả mô tả vào RAG pipeline: '[User sent an image] Image description: {mô tả lỗi}'. Kết quả: Bot trả lời đúng nguyên nhân lỗi và đề xuất giải pháp từ knowledge base. Vision description accuracy: 90%+ match với nội dung ảnh thực tế.":
        "Kịch bản 2 — Photo with Vision: Người dùng chụp ảnh màn hình lỗi phần mềm và gửi kèm caption 'Làm sao để sửa lỗi này?'. Bot: (a) Download ảnh từ Telegram server (file_id → getFile → download_file); (b) Encode base64, tạo data URL; (c) Gọi OpenRouter vision model (gpt-4o-mini) với prompt mô tả ảnh ngắn bằng tiếng Việt; (d) Feed mô tả vào RAG pipeline dưới dạng '[User sent an image] Image description: ...'. Kết quả kỳ vọng: bot kết hợp mô tả ảnh với knowledge base để đề xuất hướng xử lý, còn độ chính xác phụ thuộc chất lượng ảnh và model vision.",
    "Kịch bản 3 — Document Upload: Người dùng gửi file PDF hợp đồng bảo hiểm (2.3 MB). Bot: (a) Kiểm tra file_size ≤ 20MB (Telegram limit); (b) Download file; (c) Upload lên MinIO với path 'telegram/{bot_id}/{uuid}/{file_name}'; (d) Trích text preview (2000 bytes) cho text/plain, text/html; (e) Feed context vào RAG pipeline. Kết quả: Bot phân tích và trả lời câu hỏi dựa trên nội dung document, kèm citation đến file gốc.":
        "Kịch bản 3 — Document Upload: Người dùng gửi file qua Telegram. Bot: (a) Kiểm tra file_size ≤ 20MB; (b) Download file; (c) Upload lên MinIO với path 'telegram/{bot_id}/{uuid}/{file_name}'; (d) Với text/plain hoặc text/html, trích preview 2000 bytes làm context; với PDF/DOCX/XLSX, hiện lưu metadata/file context chứ chưa chạy full Knowledge Base ingestion. Để có structured PDF parsing và citation đầy đủ, tài liệu vẫn nên được upload qua Knowledge Base pipeline chính.",
    "•  Kênh Facebook Messenger: Cùng câu hỏi gửi qua Facebook group chat (có @mention bot). Backend gọi cùng RAG service với session_id='fb_{thread_id}_{sender_id}'. Câu trả lời format với @mention nhưng nội dung facts và citations giữ nguyên — xác nhận RAG engine hoàn toàn độc lập với channel presentation layer.\nKết quả mở rộng với Telegram và Facebook Messenger: Nội dung câu trả lời nhất quán 100% với Web và Zalo trên cả hai kênh mới, xác nhận kiến trúc channel-agnostic RAG hoạt động chính xác. Khác biệt duy nhất là latency tổng thể: Telegram ~3.2s (webhook + aiogram), Facebook Messenger ~4.1s (MQTT + worker + HMAC overhead), Zalo Bot Direct ~2.9s (HTTP trực tiếp, ít overhead nhất).: Nội dung câu trả lời nhất quán 100% với Web và Zalo, xác nhận kiến trúc channel-agnostic RAG hoạt động chính xác. Khác biệt duy nhất là latency tổng cao hơn 1-2 giây do thêm bước MQTT relay và group context fetching. Tuy nhiên, typing indicator (được worker gửi ngay khi nhận tin nhắn) giúp che giấu latency này hiệu quả.":
        "•  Kênh Facebook Messenger: Cùng câu hỏi gửi qua Facebook group chat (có @mention bot). Backend gọi cùng RAG service với session_id='fb_{thread_id}_{sender_id}'. Câu trả lời format với @mention nhưng nội dung facts và citations giữ nguyên — xác nhận RAG engine độc lập với channel presentation layer. Khác biệt chính là latency tổng cao hơn do thêm bước MQTT relay, group context fetching và HMAC HTTP giữa worker và backend; typing indicator giúp giảm cảm giác chờ của người dùng.",
    "•  Cache miss — embedding phase (t=0 → t≈400ms): Query embedding (text-embedding-3-small, OpenRouter) chiếm phần lớn latency giai đoạn đầu. Với Redis embedding cache (TTL=86400s), câu hỏi lặp lại bỏ qua giai đoạn này.":
        "•  Cache miss — embedding phase: Query embedding (text-embedding-3-small qua OpenRouter) chiếm phần latency đầu tiên. Với Redis embedding cache, câu hỏi lặp lại có thể bỏ qua giai đoạn này nếu cache key còn hiệu lực.",
    "•  Cache miss — hybrid search phase (t≈400ms → t≈1700ms): Vector search + BM25 + RRF + Cross-Encoder reranking. Cross-Encoder (ms-marco-MiniLM-L-6-v2) chiếm khoảng 400-600ms trên CPU (giảm xuống ~50ms khi có MPS/GPU). Thay bằng BAAI/bge-reranker-v2-m3 trên M1/M2 MPS cho chất lượng tốt hơn với cùng tốc độ.":
        "•  Cache miss — hybrid search phase: Qdrant prefetch dense+sparse, RRF fusion và Cross-Encoder reranking. Latency phụ thuộc OpenRouter embedding, kích thước candidate pool và CPU/MPS của reranker; benchmark isolated gần nhất đo retrieval khoảng 3.0 giây trên fixture nhỏ.",
    "•  ColPali và Multimodal RAG: ColPali (Faysse et al., ICLR 2025) và các biến thể (ColQwen2.5, ColSmolVLM) mở ra hướng tiếp cận mới cho document retrieval không cần OCR. OmniRAG đang thử nghiệm ColQwen2.5-7B để thay thế pipeline OCR→chunk→embed truyền thống, đặc biệt cho tài liệu tiếng Việt có nhiều bảng biểu và công thức. Thách thức: yêu cầu GPU (8-16GB VRAM), dung lượng storage cao hơn, và interpretability — cần bổ sung cơ chế highlight vùng ảnh liên quan thay vì chỉ trả về text chunk.":
        "•  ColPali và Multimodal RAG: ColPali (Faysse et al., ICLR 2025) và các biến thể (ColQwen2.5, ColSmolVLM) mở ra hướng tiếp cận mới cho document retrieval không cần OCR. Đây là hướng nghiên cứu tương lai cho OmniRAG, đặc biệt với tài liệu tiếng Việt có nhiều bảng biểu và công thức. Pipeline hiện tại vẫn dùng CPU-only OpenDataLoader + dense/sparse retrieval; ColPali yêu cầu GPU, storage lớn hơn và cơ chế citation/highlight vùng ảnh rõ ràng hơn trước khi đưa vào production.",
    "•  GPU-accelerated HNSW: Qdrant 2025 bổ sung GPU-accelerated HNSW indexing (order-of-magnitude faster ingestion) và 1.5-bit quantization (75% memory reduction). Upgrade Qdrant version để tận dụng khi scale lên hàng triệu vectors.":
        "•  GPU-accelerated indexing: Nếu scale lên hàng triệu vectors, có thể đánh giá các tính năng tăng tốc/quantization mới của Qdrant. Ở cấu hình hiện tại, OmniRAG pin Qdrant v1.16.0 và ưu tiên CPU-only deployment để giảm độ nặng Docker local.",
    "•  Tích hợp thực tế năm kênh giao tiếp: Zalo Bot Direct (bot-api.zapps.me), Zalo OA qua func.vn Hub, Zalo Personal Account (zca-js + isolated worker), Telegram Bot (aiogram 3 + Telegram Bot API webhook), và Facebook Messenger (fb-channel-worker + MQTT), với cơ chế HMAC webhook verification và typing indicator UX. Đây là tính năng hiếm thấy trong các RAG framework mã nguồn mở hiện tại.":
        "•  Tích hợp thực tế năm kênh giao tiếp: Zalo Bot Direct (bot-api.zapps.me), Zalo OA qua func.vn Hub, Zalo Personal Account (zca-js + isolated worker), Telegram Bot (aiogram 3 + Telegram Bot API webhook), và Facebook Messenger (fb-channel-worker + MQTT), với cơ chế webhook secret/HMAC verification tùy kênh và typing indicator UX. Đây là tính năng hiếm thấy trong các RAG framework mã nguồn mở hiện tại.",
    "•  Tích hợp Telegram Bot qua aiogram 3 và Telegram Bot API webhook, hỗ trợ text, ảnh (vision model description, gpt-4o-mini) và document upload (PDF, DOCX, TXT), với cơ chế HMAC secret token verification. Đây là kênh thứ năm được tích hợp, hoàn thiện chiến lược omnichannel với đầy đủ các nền tảng nhắn tin phổ biến tại Việt Nam. [Hình KL.1 — Telegram Bot Demo: Screenshot cuộc trò chuyện Telegram với câu hỏi thực tế và câu trả lời từ OmniRAG]":
        "•  Tích hợp Telegram Bot qua aiogram 3 và Telegram Bot API webhook, hỗ trợ text, ảnh (vision model description, gpt-4o-mini) và document upload giới hạn 20MB; text/html được trích preview, các file phức tạp được lưu metadata/file context và nên đi qua Knowledge Base upload nếu cần parsing đầy đủ. Xác thực webhook dùng secret token header của Telegram với constant-time comparison. [Hình KL.1 — Telegram Bot Demo: Screenshot cuộc trò chuyện Telegram với câu hỏi thực tế và câu trả lời từ OmniRAG]",
}


def set_paragraph_text(paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def patch_docx_text(src: Path, out: Path) -> None:
    doc = Document(src)
    found = set()
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text in REPLACEMENTS:
            set_paragraph_text(paragraph, REPLACEMENTS[text])
            found.add(text)
    missing = sorted(set(REPLACEMENTS) - found)
    if missing:
        raise SystemExit(f"Missing {len(missing)} expected paragraphs: {missing[:3]}")
    doc.save(out)


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size=size)
        except OSError:
            pass
    return ImageFont.load_default()


def wrap(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, width: int) -> list[str]:
    lines: list[str] = []
    for raw in text.split("\n"):
        words = raw.split()
        current = ""
        for word in words:
            trial = f"{current} {word}".strip()
            if draw.textbbox((0, 0), trial, font=font)[2] <= width or not current:
                current = trial
            else:
                lines.append(current)
                current = word
        lines.append(current)
    return lines


def box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], text: str, fill: str, outline: str, font, bold_font) -> None:
    draw.rounded_rectangle(xy, radius=14, fill=fill, outline=outline, width=2)
    x1, y1, x2, y2 = xy
    lines = wrap(draw, text, bold_font, x2 - x1 - 24)
    line_h = 19
    total = len(lines) * line_h
    y = y1 + (y2 - y1 - total) // 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=bold_font)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, fill="#222222", font=bold_font)
        y += line_h


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], label: str = "") -> None:
    draw.line([start, end], fill="#333333", width=3)
    ex, ey = end
    sx, sy = start
    if ex >= sx:
        pts = [(ex, ey), (ex - 12, ey - 7), (ex - 12, ey + 7)]
    else:
        pts = [(ex, ey), (ex + 12, ey - 7), (ex + 12, ey + 7)]
    draw.polygon(pts, fill="#333333")
    if label:
        font = load_font(14)
        mid = ((sx + ex) // 2, (sy + ey) // 2 - 22)
        bbox = draw.textbbox((0, 0), label, font=font)
        pad = 5
        draw.rounded_rectangle((mid[0] - (bbox[2]-bbox[0])//2 - pad, mid[1] - pad, mid[0] + (bbox[2]-bbox[0])//2 + pad, mid[1] + (bbox[3]-bbox[1]) + pad), radius=5, fill="#ffffff", outline="#eeeeee")
        draw.text((mid[0] - (bbox[2]-bbox[0])//2, mid[1]), label, fill="#444444", font=font)


def make_hybrid_png(path: Path) -> None:
    w, h = 1437, 375
    img = Image.new("RGB", (w, h), "#ffffff")
    draw = ImageDraw.Draw(img)
    title = load_font(24, True)
    body = load_font(17)
    bold = load_font(16, True)
    draw.text((w // 2 - 210, 20), "Qdrant v3 Hybrid Retrieval", fill="#222222", font=title)
    nodes = [
        ((50, 135, 220, 210), "User\nQuery", "#dae8fc", "#6c8ebf"),
        ((285, 95, 515, 175), "Dense query vector\nOpenRouter 1536-d", "#d5e8d4", "#82b366"),
        ((285, 225, 515, 305), "Sparse BM25 vector\nFastEmbed + IDF", "#ffe6cc", "#d79b00"),
        ((590, 95, 835, 175), "Qdrant prefetch\ndense top-k", "#d5e8d4", "#82b366"),
        ((590, 225, 835, 305), "Qdrant prefetch\nbm25 top-k", "#ffe6cc", "#d79b00"),
        ((910, 160, 1075, 240), "RRF fusion", "#fff2cc", "#d6b656"),
        ((1150, 160, 1335, 240), "Cross-Encoder\nrerank + top-k", "#e1d5e7", "#9673a6"),
    ]
    for xy, text, fill, outline in nodes:
        box(draw, xy, text, fill, outline, body, bold)
    arrow(draw, (220, 172), (285, 135), "embed")
    arrow(draw, (220, 172), (285, 265), "bm25")
    arrow(draw, (515, 135), (590, 135))
    arrow(draw, (515, 265), (590, 265))
    arrow(draw, (835, 135), (910, 200), "ranked")
    arrow(draw, (835, 265), (910, 200), "ranked")
    arrow(draw, (1075, 200), (1150, 200), "pool")
    draw.text((70, 330), "No Qdrant MatchText payload filter: sparse retrieval is stored as a named sparse vector.", fill="#666666", font=body)
    img.save(path, optimize=True)


def make_qdrant_png(path: Path) -> None:
    w, h = 1436, 764
    img = Image.new("RGB", (w, h), "#ffffff")
    draw = ImageDraw.Draw(img)
    title = load_font(25, True)
    body = load_font(17)
    bold = load_font(16, True)
    draw.text((w // 2 - 280, 24), "Qdrant Collection v3 Schema", fill="#222222", font=title)
    draw.rounded_rectangle((45, 80, 930, 665), radius=14, fill="#fbfdff", outline="#6c8ebf", width=2)
    draw.text((70, 105), "omnirag_openrouter_collection_v3", fill="#336699", font=load_font(20, True))
    boxes = [
        ((85, 165, 315, 245), "Point ID\nstable UUID/hash", "#f5f5f5", "#999999"),
        ((380, 145, 660, 245), "Named vector: dense\n1536-d cosine\nOpenRouter embedding", "#d5e8d4", "#82b366"),
        ((380, 285, 660, 385), "Named vector: bm25\nsparse vector\nFastEmbed + IDF", "#ffe6cc", "#d79b00"),
        ((715, 170, 885, 355), "Payload\nbot_id\ndocument_id\nsource\ntext\nparent_text\ncontext_prefix", "#dae8fc", "#6c8ebf"),
        ((85, 440, 315, 570), "Structured metadata\npage_numbers\nbboxes\nelement_types\nartifact_paths", "#e1d5e7", "#9673a6"),
        ((380, 455, 660, 555), "Payload indexes\nbot_id\ndocument_id\nsource", "#fff2cc", "#d6b656"),
        ((715, 455, 885, 555), "Retrieval\nprefetch dense + bm25\nRRF fusion", "#fff2cc", "#d6b656"),
    ]
    for xy, text, fill, outline in boxes:
        box(draw, xy, text, fill, outline, body, bold)
    arrow(draw, (315, 205), (380, 195))
    arrow(draw, (315, 205), (380, 335))
    arrow(draw, (660, 195), (715, 235))
    arrow(draw, (660, 335), (715, 285))
    arrow(draw, (315, 505), (380, 505))
    arrow(draw, (660, 505), (715, 505))
    draw.rounded_rectangle((995, 110, 1375, 590), radius=14, fill="#fcfcfc", outline="#999999", width=2)
    draw.text((1020, 135), "Payload fields used for citations", fill="#444444", font=load_font(19, True))
    bullets = [
        "document_id: exact delete/reindex scope",
        "page_numbers: page-level citation",
        "bboxes: layout-aware source region",
        "element_types: paragraph/table/image/formula",
        "heading_path: section context",
        "artifact_paths: markdown/json/images in MinIO",
    ]
    y = 190
    for item in bullets:
        for line_no, line in enumerate(wrap(draw, f"- {item}", body, 315)):
            draw.text((1030, y), line, fill="#444444", font=body)
            y += 24
        y += 12
    legacy_font = load_font(16, True)
    legacy_y = 530
    for line in wrap(draw, "Legacy TEXT MatchText index is not used for hybrid retrieval in v3.", legacy_font, 320):
        draw.text((1020, legacy_y), line, fill="#a61c1c", font=legacy_font)
        legacy_y += 24
    img.save(path, optimize=True)


def replace_docx_media(docx_path: Path, replacements: dict[str, Path]) -> None:
    tmp = docx_path.with_suffix(".tmp.docx")
    with ZipFile(docx_path, "r") as zin, ZipFile(tmp, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename in replacements:
                zout.writestr(item, replacements[item.filename].read_bytes())
            else:
                zout.writestr(item, zin.read(item.filename))
    tmp.replace(docx_path)


def main() -> int:
    if not SRC.exists():
        raise SystemExit(f"Missing source docx: {SRC}")
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    hybrid_png = ASSET_DIR / "fig_2_3_1_hybrid_search_v3_docx.png"
    qdrant_png = ASSET_DIR / "fig_2_8_qdrant_schema_v3_docx.png"
    make_hybrid_png(hybrid_png)
    make_qdrant_png(qdrant_png)
    patch_docx_text(SRC, OUT)
    replace_docx_media(
        OUT,
        {
            "word/media/image9.png": hybrid_png,
            "word/media/image14.png": qdrant_png,
        },
    )
    shutil.copy2(SRC, SRC.with_name("Decuong_OmniRAG_v4_before_accuracy_fix.docx"))
    print(f"Wrote {OUT}")
    print(f"Backup {SRC.with_name('Decuong_OmniRAG_v4_before_accuracy_fix.docx')}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
