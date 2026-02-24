"""
Script tạo đề cương OmniRAG - Phiên bản cập nhật
- Đổi tên chủ đề theo yêu cầu
- Gộp Chương 2 & 3 thành "Cơ sở lý thuyết và công nghệ sử dụng"
- Bổ sung đầy đủ thông tin từ toàn bộ repo
- Font Times New Roman, giữ nguyên format
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def set_page_border(section, color="2E598A", val="double"):
    """Thiết lập border cho trang (trang bìa)"""
    sec_pr = section._sectPr
    for pb in sec_pr.findall(qn('w:pgBorders')):
        sec_pr.remove(pb)
    pg_borders = OxmlElement('w:pgBorders')
    pg_borders.set(qn('w:offsetFrom'), 'page')
    for border_name in ['top', 'left', 'bottom', 'right']:
        edge = OxmlElement(f'w:{border_name}')
        edge.set(qn('w:val'), val) 
        edge.set(qn('w:sz'), '12')      
        edge.set(qn('w:space'), '24')
        edge.set(qn('w:color'), color)
        pg_borders.append(edge)
    sec_pr.append(pg_borders)

def remove_page_border(section):
    """Xóa border trang (trang nội dung)"""
    sec_pr = section._sectPr
    for pb in sec_pr.findall(qn('w:pgBorders')):
        sec_pr.remove(pb)
    pg_borders = OxmlElement('w:pgBorders')
    pg_borders.set(qn('w:offsetFrom'), 'page')
    for border_name in ['top', 'left', 'bottom', 'right']:
        edge = OxmlElement(f'w:{border_name}')
        edge.set(qn('w:val'), 'none')
        edge.set(qn('w:sz'), '0')
        edge.set(qn('w:space'), '0')
        pg_borders.append(edge)
    sec_pr.append(pg_borders)

def add_toc_entry(doc, text, level=0):
    """Thêm mục vào Mục lục với chấm chấm"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3 * level)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    if level == 0: 
        run.bold = True
    p.add_run('\t') 
    p.add_run('...') 

def set_times_new_roman(run, size=13, bold=False):
    """Utility: Set font Times New Roman"""
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return run

def create_de_cuong_omnirag():
    """Tạo đề cương OmniRAG hoàn chỉnh"""
    doc = Document()
    
    # ========== TÊN ĐỀ TÀI MỚI ==========
    FINAL_TITLE = "XÂY DỰNG NỀN TẢNG OMNIRAG HỖ TRỢ KHỞI TẠO TRỢ LÝ ẢO THÔNG MINH, KẾT NỐI ĐA KÊNH TỪ DỮ LIỆU TÙY CHỈNH"
    
    # ========== TRANG BÌA 1 ==========
    section = doc.sections[0]
    section.top_margin = Inches(0.79)
    section.bottom_margin = Inches(0.79)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(0.79)
    set_page_border(section, color="2E598A") 

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BỘ GIÁO DỤC VÀ ĐÀO TẠO\nTRƯỜNG ĐẠI HỌC MỎ - ĐỊA CHẤT\n---------------------------------------')
    set_times_new_roman(run, 14, bold=True)
    
    doc.add_paragraph('\n')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[HỌ VÀ TÊN SINH VIÊN]')
    set_times_new_roman(run, 14, bold=True)
    
    doc.add_paragraph('\n' * 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ĐỒ ÁN TỐT NGHIỆP\nNGÀNH CÔNG NGHỆ THÔNG TIN')
    set_times_new_roman(run, 18, bold=True)
    
    doc.add_paragraph('\n')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TÊN ĐỀ TÀI:\n')
    set_times_new_roman(run, 14, bold=True)
    run = p.add_run(FINAL_TITLE)
    set_times_new_roman(run, 17, bold=True)
    
    doc.add_paragraph('\n' * 3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Hà Nội - Năm 2026')
    set_times_new_roman(run, 12)
    
    doc.add_page_break()

    # ========== TRANG BÌA 2 ==========
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BỘ GIÁO DỤC VÀ ĐÀO TẠO\nTRƯỜNG ĐẠI HỌC MỎ - ĐỊA CHẤT\n---------------------------------------')
    set_times_new_roman(run, 14, bold=True)
    
    doc.add_paragraph('\n' * 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ĐỒ ÁN TỐT NGHIỆP\nCHUYÊN NGÀNH: KHOA HỌC MÁY TÍNH')
    set_times_new_roman(run, 18, bold=True)
    
    doc.add_paragraph('\n' * 3)
    table = doc.add_table(rows=4, cols=2)
    table.cell(0,0).text = "GIÁO VIÊN HƯỚNG DẪN:"
    table.cell(0,1).text = "ThS. [Tên Thầy/Cô]"
    table.cell(1,0).text = "SINH VIÊN THỰC HIỆN:"
    table.cell(1,1).text = "[Tên của Bro]"
    table.cell(2,0).text = "BỘ MÔN:"
    table.cell(2,1).text = "KHOA HỌC MÁY TÍNH"
    table.cell(3,0).text = "LỚP:"
    table.cell(3,1).text = "KHMT ỨNG DỤNG [Lớp]"
    
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for r in para.runs:
                    set_times_new_roman(r, 13, bold=True)
    
    doc.add_paragraph('\n' * 8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Hà Nội – Năm 2026')
    set_times_new_roman(run, 12)

    # ========== MỤC LỤC ==========
    new_sec = doc.add_section()
    remove_page_border(new_sec)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('MỤC LỤC')
    set_times_new_roman(run, 16, bold=True)
    
    # Phần mở đầu
    add_toc_entry(doc, "LỜI CẢM ƠN", 0)
    add_toc_entry(doc, "DANH MỤC KÝ HIỆU VÀ CHỮ VIẾT TẮT", 0)
    add_toc_entry(doc, "DANH MỤC BẢNG BIỂU", 0)
    add_toc_entry(doc, "DANH MỤC HÌNH ẢNH", 0)
    add_toc_entry(doc, "MỞ ĐẦU", 0)
    add_toc_entry(doc, "1. Lý do chọn đề tài", 1)
    add_toc_entry(doc, "2. Mục tiêu của đề tài", 1)
    add_toc_entry(doc, "3. Đối tượng và phạm vi nghiên cứu", 1)
    add_toc_entry(doc, "4. Phương pháp nghiên cứu", 1)
    add_toc_entry(doc, "5. Bố cục đề tài", 1)
    
    # ========== CHƯƠNG 1: TỔNG QUAN ==========
    add_toc_entry(doc, "CHƯƠNG 1. TỔNG QUAN VÀ KHẢO SÁT HỆ THỐNG", 0)
    add_toc_entry(doc, "1.1. Giới thiệu bài toán trợ lý ảo thông minh từ dữ liệu tùy chỉnh", 1)
    add_toc_entry(doc, "1.2. Nhu cầu giao tiếp đa kênh (Omnichannel) trong doanh nghiệp hiện đại", 1)
    add_toc_entry(doc, "1.3. Thách thức trong triển khai trợ lý ảo đa nền tảng", 1)
    add_toc_entry(doc, "1.3.1. Vấn đề bảo mật và tính riêng tư dữ liệu", 2)
    add_toc_entry(doc, "1.3.2. Khó khăn tích hợp đa nguồn dữ liệu (PDF, DOCX, Website, API)", 2)
    add_toc_entry(doc, "1.3.3. Chi phí cao và phức tạp kỹ thuật", 2)
    add_toc_entry(doc, "1.4. Khảo sát các giải pháp hiện có", 1)
    add_toc_entry(doc, "1.4.1. Các nền tảng chatbot đóng gói sẵn (Dialogflow, Rasa, Botpress)", 2)
    add_toc_entry(doc, "1.4.2. Đánh giá khả năng tùy chỉnh và tích hợp RAG của các nền tảng", 2)
    add_toc_entry(doc, "1.4.3. So sánh giải pháp mã nguồn đóng vs mã nguồn mở", 2)
    add_toc_entry(doc, "1.5. Phân tích yêu cầu hệ thống OmniRAG", 1)
    add_toc_entry(doc, "1.5.1. Yêu cầu chức năng: Self-service Bot Builder, Multi-tenancy, Omnichannel", 2)
    add_toc_entry(doc, "1.5.2. Yêu cầu phi chức năng: Hiệu năng, Bảo mật, Khả năng mở rộng", 2)
    add_toc_entry(doc, "1.5.3. Yêu cầu tích hợp: OpenRouter, Qdrant, MinIO, Redis, Docker", 2)
    
    # ========== CHƯƠNG 2: CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ (GỘP) ==========
    add_toc_entry(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ SỬ DỤNG", 0)
    
    # Phần 2.1: LÝ THUYẾT AI VÀ RAG
    add_toc_entry(doc, "2.1. Cơ sở lý thuyết về AI và Generative AI", 1)
    add_toc_entry(doc, "2.1.1. Tổng quan về Large Language Models (LLM)", 2)
    add_toc_entry(doc, "2.1.2. Kiến trúc Transformer và cơ chế Attention", 2)
    add_toc_entry(doc, "2.1.3. Fine-tuning vs Prompt Engineering vs RAG", 2)
    
    add_toc_entry(doc, "2.2. Kiến trúc Retrieval-Augmented Generation (RAG)", 1)
    add_toc_entry(doc, "2.2.1. Quy trình RAG cơ bản: Index - Retrieve - Generate", 2)
    add_toc_entry(doc, "2.2.2. Phân loại RAG: Naive RAG vs Advanced RAG vs Modular RAG", 2)
    add_toc_entry(doc, "2.2.3. Vector Embeddings và Semantic Search", 2)
    
    add_toc_entry(doc, "2.3. Các kỹ thuật tối ưu hóa RAG nâng cao", 1)
    add_toc_entry(doc, "2.3.1. Hybrid Search: Kết hợp Dense (Vector) và Sparse (BM25) Retrieval", 2)
    add_toc_entry(doc, "2.3.2. Query Transformation: HyDE và Multi-Query Generation", 2)
    add_toc_entry(doc, "2.3.3. Document Re-ranking với Cross-Encoders", 2)
    add_toc_entry(doc, "2.3.4. Chiến lược Chunking: Recursive vs Semantic Splitting", 2)
    add_toc_entry(doc, "2.3.5. Conversation Memory và Context Window Management", 2)
    
    # Phần 2.4: CÔNG NGHỆ Backend & Infrastructure
    add_toc_entry(doc, "2.4. Công nghệ Backend và hạ tầng hệ thống", 1)
    add_toc_entry(doc, "2.4.1. FastAPI: Framework web bất đồng bộ hiệu năng cao", 2)
    add_toc_entry(doc, "2.4.2. Python async/await và ASGI servers (Uvicorn)", 2)
    add_toc_entry(doc, "2.4.3. SQLAlchemy ORM và Alembic migrations", 2)
    add_toc_entry(doc, "2.4.4. Golang API Gateway: Gin framework và middleware pattern", 2)
    
    add_toc_entry(doc, "2.5. Công nghệ AI và Machine Learning", 1)
    add_toc_entry(doc, "2.5.1. LangChain: Xây dựng Chains và tích hợp Tools", 2)
    add_toc_entry(doc, "2.5.2. OpenRouter: Unified API cho 400+ AI models", 2)
    add_toc_entry(doc, "2.5.3. Sentence Transformers và Embedding Models", 2)
    add_toc_entry(doc, "2.5.4. Spacy: Xử lý ngôn ngữ tự nhiên (NLP)", 2)
    
    add_toc_entry(doc, "2.6. Cơ sở dữ liệu và lưu trữ", 1)
    add_toc_entry(doc, "2.6.1. Qdrant Vector Database và HNSW Indexing", 2)
    add_toc_entry(doc, "2.6.2. PostgreSQL: Quan hệ dữ liệu Users, Tenants, Bots, Documents", 2)
    add_toc_entry(doc, "2.6.3. MongoDB: Lưu trữ chat logs và analytics", 2)
    add_toc_entry(doc, "2.6.4. Redis: Caching, Session Management và Celery Broker", 2)
    add_toc_entry(doc, "2.6.5. MinIO: Object Storage tương thích S3 cho file uploads", 2)
    
    add_toc_entry(doc, "2.7. Công nghệ Frontend và UX", 1)
    add_toc_entry(doc, "2.7.1. React 19: Component-based architecture và Hooks", 2)
    add_toc_entry(doc, "2.7.2. TypeScript: Type-safe development", 2)
    add_toc_entry(doc, "2.7.3. Vite: Lightning-fast build tool", 2)
    add_toc_entry(doc, "2.7.4. Tailwind CSS: Utility-first CSS framework", 2)
    add_toc_entry(doc, "2.7.5. Zustand: State management cho React", 2)
    
    add_toc_entry(doc, "2.8. DevOps và Microservices", 1)
    add_toc_entry(doc, "2.8.1. Docker và Docker Compose: Container orchestration", 2)
    add_toc_entry(doc, "2.8.2. Celery: Distributed task queue cho xử lý bất đồng bộ", 2)
    add_toc_entry(doc, "2.8.3. Multi-stage Docker builds và optimization", 2)
    add_toc_entry(doc, "2.8.4. Rate Limiting và API Gateway Pattern", 2)
    
    add_toc_entry(doc, "2.9. Công nghệ tích hợp đa kênh (Omnichannel)", 1)
    add_toc_entry(doc, "2.9.1. Telegram Bot API: Webhook và Long Polling", 2)
    add_toc_entry(doc, "2.9.2. Zalo OA và Facebook Messenger integration", 2)
    add_toc_entry(doc, "2.9.3. Web Widget: Embeddable chat component", 2)
    add_toc_entry(doc, "2.9.4. REST API và WebSocket cho real-time communication", 2)
    
    # ========== CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ ==========
    add_toc_entry(doc, "CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ NỀN TẢNG OMNIRAG", 0)
    
    add_toc_entry(doc, "3.1. Kiến trúc tổng thể hệ thống", 1)
    add_toc_entry(doc, "3.1.1. Kiến trúc Microservices: Frontend - Gateway - Backend - Services", 2)
    add_toc_entry(doc, "3.1.2. Luồng xử lý request từ Client đến AI Response", 2)
    add_toc_entry(doc, "3.1.3. Multi-tenancy architecture với data isolation", 2)
    
    add_toc_entry(doc, "3.2. Phân tích chức năng và luồng nghiệp vụ", 1)
    add_toc_entry(doc, "3.2.1. Quy trình đăng ký và xác thực (JWT Authentication)", 2)
    add_toc_entry(doc, "3.2.2. Quy trình tạo Bot và upload documents", 2)
    add_toc_entry(doc, "3.2.3. Document Processing Pipeline: Upload → Parse → Chunk → Embed → Index", 2)
    add_toc_entry(doc, "3.2.4. RAG Chat Flow: Query → Transform → Retrieve → Rank → Generate", 2)
    add_toc_entry(doc, "3.2.5. Omnichannel Message Routing với Channel Gateway", 2)
    
    add_toc_entry(doc, "3.3. Thiết kế cơ sở dữ liệu", 1)
    add_toc_entry(doc, "3.3.1. ERD cho PostgreSQL: Users, Tenants, Bots, Documents", 2)
    add_toc_entry(doc, "3.3.2. Schema MongoDB: Conversations, Sessions, Chat Logs", 2)
    add_toc_entry(doc, "3.3.3. Qdrant Collection structure: Embeddings và Metadata", 2)
    add_toc_entry(doc, "3.3.4. Redis key design cho caching và rate limiting", 2)
    
    add_toc_entry(doc, "3.4. Thiết kế các module chính", 1)
    add_toc_entry(doc, "3.4.1. Auth Module: User registration, login, JWT", 2)
    add_toc_entry(doc, "3.4.2. Bot Management Module: CRUD operations", 2)
    add_toc_entry(doc, "3.4.3. Document Processing Module: Celery tasks, chunking strategies", 2)
    add_toc_entry(doc, "3.4.4. RAG Service: OpenRouter integration, hybrid search, re-ranking", 2)
    add_toc_entry(doc, "3.4.5. Channel Connectors: Telegram, Zalo, Web Widget adapters", 2)
    add_toc_entry(doc, "3.4.6. Golang Gateway: Proxy, caching, rate limiting", 2)
    
    add_toc_entry(doc, "3.5. Thiết kế giao diện (UI/UX)", 1)
    add_toc_entry(doc, "3.5.1. Dashboard: Thống kê và quản lý Bots", 2)
    add_toc_entry(doc, "3.5.2. Bot Builder: Upload documents và cấu hình", 2)
    add_toc_entry(doc, "3.5.3. Chat Playground: Test Bot với conversation history", 2)
    add_toc_entry(doc, "3.5.4. Channel Configuration: Setup Telegram, Zalo, Web Widget", 2)
    add_toc_entry(doc, "3.5.5. Analytics Dashboard: Chat logs và performance metrics", 2)
    
    add_toc_entry(doc, "3.6. Các biểu đồ hệ thống", 1)
    add_toc_entry(doc, "3.6.1. Use Case Diagram: Actors và chức năng chính", 2)
    add_toc_entry(doc, "3.6.2. Sequence Diagram: RAG Chat workflow", 2)
    add_toc_entry(doc, "3.6.3. Activity Diagram: Document processing pipeline", 2)
    add_toc_entry(doc, "3.6.4. Component Diagram: Microservices architecture", 2)
    add_toc_entry(doc, "3.6.5. Deployment Diagram: Docker Compose orchestration", 2)
    
    # ========== CHƯƠNG 4: CÀI ĐẶT VÀ TRIỂN KHAI ==========
    add_toc_entry(doc, "CHƯƠNG 4. CÀI ĐẶT VÀ TRIỂN KHAI HỆ THỐNG", 0)
    
    add_toc_entry(doc, "4.1. Môi trường phát triển và công cụ", 1)
    add_toc_entry(doc, "4.1.1. Setup Docker và Docker Compose", 2)
    add_toc_entry(doc, "4.1.2. Cấu hình Python virtual environment và dependencies", 2)
    add_toc_entry(doc, "4.1.3. Setup Node.js và React development environment", 2)
    add_toc_entry(doc, "4.1.4. Cấu hình Golang module và dependencies", 2)
    
    add_toc_entry(doc, "4.2. Triển khai Backend services", 1)
    add_toc_entry(doc, "4.2.1. Cài đặt FastAPI với Uvicorn workers", 2)
    add_toc_entry(doc, "4.2.2. Setup PostgreSQL với Alembic migrations", 2)
    add_toc_entry(doc, "4.2.3. Cấu hình Qdrant vector database", 2)
    add_toc_entry(doc, "4.2.4. Setup MinIO object storage", 2)
    add_toc_entry(doc, "4.2.5. Cấu hình Redis cho caching và Celery broker", 2)
    add_toc_entry(doc, "4.2.6. Triển khai Celery workers cho document processing", 2)
    
    add_toc_entry(doc, "4.3. Triển khai Golang API Gateway", 1)
    add_toc_entry(doc, "4.3.1. Setup Gin router và middleware stack", 2)
    add_toc_entry(doc, "4.3.2. Implement proxy handler với caching", 2)
    add_toc_entry(doc, "4.3.3. Rate limiting với Redis counters", 2)
    add_toc_entry(doc, "4.3.4. Logging và monitoring với Zap", 2)
    add_toc_entry(doc, "4.3.5. Health checks và graceful shutdown", 2)
    
    add_toc_entry(doc, "4.4. Triển khai Frontend application", 1)
    add_toc_entry(doc, "4.4.1. Setup React với Vite và TypeScript", 2)
    add_toc_entry(doc, "4.4.2. Implement authentication với JWT", 2)
    add_toc_entry(doc, "4.4.3. API client với Axios và interceptors", 2)
    add_toc_entry(doc, "4.4.4. State management với Zustand", 2)
    add_toc_entry(doc, "4.4.5. Styling với Tailwind CSS và dark mode", 2)
    
    add_toc_entry(doc, "4.5. Tích hợp RAG và OpenRouter", 1)
    add_toc_entry(doc, "4.5.1. Setup OpenRouter API client", 2)
    add_toc_entry(doc, "4.5.2. Implement hybrid search với Qdrant", 2)
    add_toc_entry(doc, "4.5.3. Query transformation: HyDE và multi-query", 2)
    add_toc_entry(doc, "4.5.4. Document re-ranking pipeline", 2)
    add_toc_entry(doc, "4.5.5. Conversation memory management", 2)
    
    add_toc_entry(doc, "4.6. Tích hợp đa kênh (Omnichannel)", 1)
    add_toc_entry(doc, "4.6.1. Setup Telegram Bot với webhook", 2)
    add_toc_entry(doc, "4.6.2. Tích hợp Zalo OA qua func.vn gateway", 2)
    add_toc_entry(doc, "4.6.3. Phát triển Web Widget embeddable component", 2)
    add_toc_entry(doc, "4.6.4. Unified message adapter pattern", 2)
    
    # ========== CHƯ��NG 5: THỰC NGHIỆM VÀ ĐÁNH GIÁ ==========
    add_toc_entry(doc, "CHƯƠNG 5. THỰC NGHIỆM VÀ ĐÁNH GIÁ KẾT QUẢ", 0)
    
    add_toc_entry(doc, "5.1. Môi trường thử nghiệm", 1)
    add_toc_entry(doc, "5.1.1. Cấu hình máy chủ và Docker resources", 2)
    add_toc_entry(doc, "5.1.2. Dataset thử nghiệm: PDFs, Documents, URLs", 2)
    add_toc_entry(doc, "5.1.3. Test cases và kịch bản kiểm thử", 2)
    
    add_toc_entry(doc, "5.2. Kết quả demo và chức năng", 1)
    add_toc_entry(doc, "5.2.1. Demo Web Dashboard: Quản lý Bots và Documents", 2)
    add_toc_entry(doc, "5.2.2. Demo Chat với RAG: Truy xuất context từ documents", 2)
    add_toc_entry(doc, "5.2.3. Demo Telegram Bot: Chat đa kênh real-time", 2)
    add_toc_entry(doc, "5.2.4. Demo Web Widget: Nhúng trên website", 2)
    add_toc_entry(doc, "5.2.5. Demo Multi-tenancy: Data isolation giữa các tenant", 2)
    
    add_toc_entry(doc, "5.3. Đánh giá hiệu năng hệ thống", 1)
    add_toc_entry(doc, "5.3.1. Latency: Thời gian response với/không cache", 2)
    add_toc_entry(doc, "5.3.2. Throughput: Requests/second với Golang Gateway", 2)
    add_toc_entry(doc, "5.3.3. Memory usage: So sánh Python vs Golang gateway", 2)
    add_toc_entry(doc, "5.3.4. Document processing speed: Upload đến indexing", 2)
    
    add_toc_entry(doc, "5.4. Đánh giá chất lượng RAG", 1)
    add_toc_entry(doc, "5.4.1. Relevance: Độ liên quan của retrieved documents", 2)
    add_toc_entry(doc, "5.4.2. Accuracy: Độ chính xác câu trả lời so với ground truth", 2)
    add_toc_entry(doc, "5.4.3. Context utilization: Sử dụng thông tin từ documents", 2)
    add_toc_entry(doc, "5.4.4. So sánh Hybrid Search vs Vector-only search", 2)
    add_toc_entry(doc, "5.4.5. Impact của re-ranking trên chất lượng", 2)
    
    add_toc_entry(doc, "5.5. Đánh giá tính nhất quán đa kênh", 1)
    add_toc_entry(doc, "5.5.1. Test cross-channel consistency: Same bot trên Telegram, Web, Zalo", 2)
    add_toc_entry(doc, "5.5.2. Conversation continuity giữa các kênh", 2)
    add_toc_entry(doc, "5.5.3. Response time comparison giữa các channels", 2)
    
    add_toc_entry(doc, "5.6. Thảo luận và phân tích", 1)
    add_toc_entry(doc, "5.6.1. Ưu điểm của kiến trúc Microservices", 2)
    add_toc_entry(doc, "5.6.2. Lợi ích của Golang Gateway cho performance", 2)
    add_toc_entry(doc, "5.6.3. Hiệu quả của advanced RAG techniques", 2)
    add_toc_entry(doc, "5.6.4. Những thách thức và hạn chế còn tồn tại", 2)
    
    # ========== PHỤ LỤC VÀ TÀI LIỆU THAM KHẢO ==========
    add_toc_entry(doc, "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 0)
    add_toc_entry(doc, "1. Kết luận", 1)
    add_toc_entry(doc, "2. Hướng phát triển tương lai", 1)
    add_toc_entry(doc, "2.1. Hỗ trợ thêm AI models (Claude, Gemini, Local LLMs)", 2)
    add_toc_entry(doc, "2.2. Agentic RAG với LangGraph", 2)
    add_toc_entry(doc, "2.3. Fine-tuning embeddings cho domain-specific", 2)
    add_toc_entry(doc, "2.4. Advanced analytics và observability với LangFuse", 2)
    add_toc_entry(doc, "2.5. Voice support và multimodal RAG", 2)
    
    add_toc_entry(doc, "DANH MỤC TÀI LIỆU THAM KHẢO", 0)
    
    add_toc_entry(doc, "PHỤ LỤC", 0)
    add_toc_entry(doc, "Phụ lục A: API Endpoints và Postman Collection", 1)
    add_toc_entry(doc, "Phụ lục B: Database Schema và ERD diagrams", 1)
    add_toc_entry(doc, "Phụ lục C: Docker Compose configuration", 1)
    add_toc_entry(doc, "Phụ lục D: Environment variables và configuration", 1)
    add_toc_entry(doc, "Phụ lục E: Code snippets: Telegram Bot, Web Widget", 1)
    add_toc_entry(doc, "Phụ lục F: Performance benchmarks và metrics", 1)
    add_toc_entry(doc, "Phụ lục G: User manual và deployment guide", 1)

    doc.add_page_break()

    # ========== NỘI DUNG MẪU ==========
    p = doc.add_paragraph()
    run = p.add_run('MỞ ĐẦU')
    set_times_new_roman(run, 16, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    p = doc.add_paragraph(
        "Nội dung chi tiết sẽ được bổ sung tại đây. Đề cương này đã được cập nhật với:\n"
        "• Tiêu đề mới: 'Xây dựng nền tảng OmniRAG hỗ trợ khởi tạo trợ lý ảo thông minh, kết nối đa kênh từ dữ liệu tùy chỉnh'\n"
        "• Chương 2 và 3 đã được gộp thành 'Cơ sở lý thuyết và công nghệ sử dụng'\n"
        "• Bổ sung đầy đủ thông tin từ repository: Golang Gateway, OpenRouter, Advanced RAG, Docker Microservices\n"
        "• Font Times New Roman được áp dụng nhất quán trong toàn bộ tài liệu"
    )
    for run in p.runs:
        set_times_new_roman(run, 13)
    
    # Lưu file
    output_file = 'OmniRAG_DeCuong_Updated.docx'
    doc.save(output_file)
    print(f"✅ HOÀN THÀNH: {output_file}")
    print(f"📋 Đã tạo đề cương với:")
    print(f"   • Tiêu đề mới: Đã cập nhật ✓")
    print(f"   • Gộp Chương 2 & 3 → Chương 2 mới ✓")
    print(f"   • Bổ sung thông tin từ repo ✓")
    print(f"   • Font Times New Roman ✓")

if __name__ == "__main__":
    create_de_cuong_omnirag()